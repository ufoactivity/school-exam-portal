import streamlit as st
import pandas as pd
import numpy as np
import io
import pulp
import traceback
import random
import openpyxl
import re
from datetime import datetime
from collections import defaultdict

# ==========================================
# 📌 套件檢查與熱修復
# ==========================================
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl.drawing.text
    openpyxl.drawing.text.Font.pitchFamily.max = 1000
except:
    pass

# ==========================================
# 1. 網頁頁面配置與全局狀態記憶體初始化
# ==========================================
st.set_page_config(page_title="段考試務全能系統", page_icon="🏫", layout="wide")
st.title("🏫 試務組 - 段考試務全能系統 (旗艦整合版)")
st.info("💡 終極升級：已將「命題出題排定」、「催繳通知單」與「段考監考排班」三大系統完美融合！一站式完成所有段考考務工作。")

# --- 狀態記憶體初始化 (Session State) ---
if 'uploader_key' not in st.session_state:
    st.session_state['uploader_key'] = 0

# 【階段一：命題出題】記憶體
if 'results_p1' not in st.session_state:
    st.session_state['results_p1'] = None
if 'debug_log_p1' not in st.session_state:
    st.session_state['debug_log_p1'] = []

# 【階段二：催繳通知】記憶體
if 'docx_data_p2_print' not in st.session_state:
    st.session_state['docx_data_p2_print'] = None
if 'docx_data_p2_msg' not in st.session_state:
    st.session_state['docx_data_p2_msg'] = None
if 'processed_p2' not in st.session_state:
    st.session_state['processed_p2'] = False

# 【階段三：監考排班】記憶體
if 'results_p3' not in st.session_state:
    st.session_state['results_p3'] = None
if 'time_rules' not in st.session_state:
    st.session_state.time_rules = pd.DataFrame([{"老師": None, "允許日期": "無限制", "允許節次": ""} for _ in range(3)])
if 'bind_rules' not in st.session_state:
    st.session_state.bind_rules = pd.DataFrame([{"老師": None, "班級": None}] * 3)
if 'last_bind_file' not in st.session_state:
    st.session_state.last_bind_file = None

# ==========================================
# 2. 輔助功能定義 (Domain Knowledge)
# ==========================================

# ----- 【共用與階段三輔助】 -----
def to_excel_bytes(df, header_df=None):
    output = io.BytesIO()
    if header_df is not None:
        df.columns = header_df.columns
        final_out = pd.concat([header_df, df], ignore_index=True)
    else:
        final_out = df
    final_out = final_out.fillna("")
    for col in final_out.columns:
        final_out[col] = final_out[col].apply(lambda x: f" {x}" if isinstance(x, str) and (x.startswith("=") or x.startswith("-")) else x)
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        final_out.to_excel(writer, index=False, header=False)
    return output.getvalue()

def normalize_cls(c):
    if pd.isna(c) or c is None: return ""
    s = str(c).strip().replace('ㄧ', '一').replace(' ', '').replace(' ', '')
    s = s.translate(str.maketrans('１２３４５６７８９０', '1234567890'))
    return s

def clean_str(s):
    if pd.isna(s) or s is None: return ""
    s = str(s).strip().replace('ㄧ', '一').replace(' ', '').replace(' ', '').replace('\n', '').replace('\r', '')
    s = s.translate(str.maketrans('１２３４５６７８９０', '1234567890'))
    return s

def normalize_subject_p3(s):
    s = clean_str(s)
    aliases = {'國文':'國語文', '英文':'英語文', '公社':'公民與社會', '公民':'公民與社會', 
               '地科':'地球科學', '健護':'健康與護理', '護理':'健康與護理', 
               '國防':'全民國防教育', '生科':'生活科技', '應數':'應用數學'}
    return aliases.get(s, s)

def get_teacher_fuzzy(cls, subj, course_dict):
    if (cls, subj) in course_dict: return course_dict[(cls, subj)]
    clean_target = subj.replace('選修', '').replace('彈性學習', '').replace('補強', '').replace('-', '')
    for (c, s), t in course_dict.items():
        if c == cls:
            s_clean = s.replace('選修', '').replace('彈性學習', '').replace('補強', '').replace('-', '')
            if clean_target and (clean_target in s_clean or s_clean in clean_target):
                return t
    return ""

def extract_mm_dd(text, default_month="05"):
    if pd.isna(text) or text is None: return ""
    s = str(text).strip().replace(' ', '')
    m = re.search(r'(\d{1,2})[-/月](\d{1,2})', s)
    if m: return f"{int(m.group(1)):02d}-{int(m.group(2)):02d}"
    m2 = re.search(r'(\d{1,2})日', s)
    if m2: return f"{int(default_month):02d}-{int(m2.group(1)):02d}"
    return ""

def get_ai_date_str(j, day_starts, ai_date_strs):
    day_idx = 0
    for idx, ds in enumerate(day_starts):
        if j >= ds: day_idx = idx
    return ai_date_strs[min(day_idx, len(ai_date_strs)-1)]

def extract_period_num(s):
    if pd.isna(s): return -1
    s = str(s).strip()
    if any(k in s for k in ['月', '日', '年', '表', '華南', '期中', '次數', '日期']): return -1
    cn_to_num = {'一':'1', '二':'2', '三':'3', '四':'4', '五':'5', 
                 '六':'6', '七':'7', '八':'8', '九':'9', '十':'10', 
                 '１':'1', '２':'2', '３':'3', '４':'4', '５':'5', '６':'6', '７':'7'}
    for k, v in cn_to_num.items(): s = s.replace(k, v)
    nums = re.findall(r'\d+', s)
    if nums:
        p = int(nums[0])
        if 1 <= p <= 15: return p
    return -1

def matches_date(val_str, d_date):
    if not d_date: return False
    s = str(val_str).replace(' ', '').replace('nan', '')
    if not s: return False
    d1 = d_date.strftime('%Y-%m-%d')
    d2 = d_date.strftime('%m-%d')
    d3 = d_date.strftime('%Y/%m/%d')
    d4 = d_date.strftime('%m月%d日')
    d5 = d_date.strftime('%d日')
    d6 = str(d_date.day) + '日'
    return any(x in s for x in [d1, d2, d3, d4, d5, d6])

# ----- 【階段一：命題出題輔助】 -----
def clean_subject_name_p1(subj_raw):
    if pd.isna(subj_raw) or subj_raw is None: return ""
    s = str(subj_raw).strip()
    s = re.sub(r'[（\(].*?[）\)]', '', s)
    s = s.replace(' ', '').replace(' ', '').replace('\n', '').replace('-', '')
    if s == '英文': return '英語文'
    if s == '國文': return '國語文'
    if s in ['數學I', '數學III', '數學V']: return '數學'
    return s.strip()

def extract_target_depts(subj_raw):
    m = re.search(r'[（\(](.*?)[）\)]', str(subj_raw))
    if not m: return [] 
    content = m.group(1) if m.group(1) else m.group(2)
    depts = []
    vocational_chars = ['商', '國', '航', '電', '資', '廣', '美', '應', '日', '觀']
    general_chars = ['高', '普', 'H'] 
    is_english = any(k in str(subj_raw) for k in ['英文', '英語'])
    if '職' in content: 
        depts.extend(vocational_chars)
        if is_english and '應' in depts: depts.remove('應')
    if '普' in content or 'H' in content: 
        depts.extend(general_chars)
        if is_english and '應' not in depts: depts.append('應')
    for char in content:
        if char in vocational_chars + general_chars and char not in depts: depts.append(char)
    if is_english:
        has_voc = any(c in depts for c in vocational_chars if c != '應')
        has_gen = any(c in depts for c in general_chars)
        if has_voc and not has_gen:
            for v in vocational_chars:
                if v != '應' and v not in depts: depts.append(v)
        if has_gen:
            if '應' not in depts: depts.append('應')
    return list(set(depts))

def build_class_teacher_dict(df_peike):
    t_dict = {}
    if df_peike.empty: return t_dict
    for _, row in df_peike.iterrows():
        subj_raw = row.iloc[0]
        if pd.isna(subj_raw) or str(subj_raw).strip() == "": continue
        subj_clean = clean_subject_name_p1(subj_raw)
        for col_name in df_peike.columns[1:]:
            if 'Unnamed' in str(col_name): break
            val = row[col_name]
            if pd.isna(val) or str(val).strip() == "": continue
            c_name = str(col_name).strip() 
            if not c_name: continue
            dept_char = c_name[0] 
            t_str = str(val).strip()
            teachers = [t.strip() for t in re.split(r'[、,，\s/]+', t_str) if t.strip()]
            key = (dept_char, subj_clean)
            if key not in t_dict: t_dict[key] = []
            for t in teachers: t_dict[key].append(t)
    return t_dict

def get_teachers_for_subject_p1(subj_raw, t_dict):
    target_depts = extract_target_depts(subj_raw)
    subj_clean = clean_subject_name_p1(subj_raw)
    teachers = []
    if not target_depts:
        for (dept, s), ts in t_dict.items():
            if s == subj_clean: teachers.extend(ts)
    else:
        mapped_prefixes = set()
        for dept in target_depts:
            if dept in ['普', 'H']: mapped_prefixes.add('高')
            else: mapped_prefixes.add(dept)
        for dept in mapped_prefixes:
            if (dept, subj_clean) in t_dict: teachers.extend(t_dict[(dept, subj_clean)])
    if not teachers:
        mapped_target = set()
        for d in target_depts:
            if d in ['普', 'H']: mapped_target.add('高')
            else: mapped_target.add(d)
        for (dept, s), ts in t_dict.items():
            if not target_depts or dept in mapped_target:
                if subj_clean in s or s in subj_clean:
                    teachers.extend(ts)
    return teachers

def parse_sync_file(df_sync):
    group_map = {}
    for r in range(len(df_sync)):
        for c in range(len(df_sync.columns)):
            val = str(df_sync.iloc[r, c]).strip()
            if val and val != 'nan':
                val = val.replace(' ', '')
                m = re.search(r'([123一二三])(?:年級)?(.*)', val)
                if m:
                    g_num = 1 if m.group(1) in ['1','一'] else 2 if m.group(1) in ['2','二'] else 3 if m.group(1) in ['3','三'] else None
                    if g_num: group_map[(g_num, clean_subject_name_p1(m.group(2)))] = r
    return group_map

def extract_history(file_history):
    history_map = {} 
    if not file_history: return history_map
    try:
        wb = openpyxl.load_workbook(file_history, data_only=True)
        ws = wb.active 
        header_row = -1
        grade_subj_cols = []
        for r in range(1, 15):
            current_row_cols = []
            for c in range(1, 25):
                if str(ws.cell(row=r, column=c).value).strip().replace(' ', '').replace(' ', '') == "科目":
                    current_row_cols.append(c)
            if len(current_row_cols) >= 3:
                header_row = r; grade_subj_cols = current_row_cols[:3]; break
        if header_row != -1 and len(grade_subj_cols) >= 3:
            grade_mapping = [{'grade': 3, 'col': grade_subj_cols[0]}, {'grade': 2, 'col': grade_subj_cols[1]}, {'grade': 1, 'col': grade_subj_cols[2]}]
            for r in range(header_row + 1, ws.max_row + 1):
                for mapping in grade_mapping:
                    subj_col = mapping['col']
                    subj_raw = ws.cell(row=r, column=subj_col).value
                    if not subj_raw or str(subj_raw).strip() == "": continue
                    past_teachers = []
                    for offset in range(1, 8):
                        target_col = subj_col + offset
                        if mapping != grade_mapping[-1] and target_col >= grade_mapping[grade_mapping.index(mapping)+1]['col']: break
                        if mapping == grade_mapping[-1] and target_col >= grade_subj_cols[-1] + 6: break
                        cell_val = ws.cell(row=r, column=target_col).value
                        if cell_val and str(cell_val).strip() not in ["", "None", "不列入", "無"]:
                            past_teachers.append(str(cell_val).strip())
                    key = (mapping['grade'], str(subj_raw).strip())
                    if key not in history_map: history_map[key] = []
                    history_map[key].extend(past_teachers)
    except: pass
    return history_map

def generate_perfect_balanced_sequence(pool, global_counts, sequence_length=10):
    if not pool: return [""] * sequence_length
    best_seq = None
    best_penalty = float('inf')
    pool_counts = {}
    for t in pool: pool_counts[t] = pool_counts.get(t, 0) + 1
    for _ in range(200):
        shuffled = pool.copy()
        random.shuffle(shuffled)
        seq = [shuffled[i % len(shuffled)] for i in range(sequence_length)]
        penalty = 0
        for i, t in enumerate(seq): penalty += global_counts[t][i] ** 2 
        first_half, second_half = seq[:5], seq[5:10]
        for t, total_c in pool_counts.items():
            if total_c > 1:
                diff = abs(first_half.count(t) - second_half.count(t))
                if diff > 1: penalty += (diff * 200)
        for i in range(len(seq) - 1):
            if seq[i] == seq[i+1] and seq[i] != "": penalty += 50
        if penalty < best_penalty:
            best_penalty = penalty
            best_seq = seq
    for i, t in enumerate(best_seq): global_counts[t][i] += 1
    return best_seq

# ==========================================
# 3. 核心介面佈局：利用 Tabs 分割三個階段
# ==========================================
tab1, tab2, tab3 = st.tabs(["🎯 階段一：命題出題教師排定", "📑 階段二：試卷催繳通知單", "📅 階段三：段考監考智能排班"])

# ---------------------------------------------------------
# 【階段一：命題出題教師排定】
# ---------------------------------------------------------
with tab1:
    st.subheader("🎯 階段一：命題與出題教師自動排定系統")
    st.markdown("自動比對配課表，智慧解析職科與普高，並支援跨學期歷史防撞與均分演算法。")
    
    col1_p1, col2_p1 = st.columns([1, 1], gap="large")

    with col1_p1:
        st.markdown("##### 📂 1. 上傳基礎資料")
        file_peike_p1 = st.file_uploader("1️⃣ 上傳配課表 (需含：一、二、三年級)", type=['xlsx'], key=f"p1_peike_{st.session_state['uploader_key']}")
        file_template_p1 = st.file_uploader("2️⃣ 上傳進度及出題總表 (空白範本)", type=['xlsx'], key=f"p1_temp_{st.session_state['uploader_key']}")
        st.write("---")
        file_sync_p1 = st.file_uploader("3️⃣ 上傳同卷設定表 (非必填)", type=['xlsx', 'csv'], key=f"p1_sync_{st.session_state['uploader_key']}")
        st.write("---")
        file_history_p1 = st.file_uploader("🕰️ 4️⃣ 上傳上學期出題總表 (下學期排班專用，確保次數延續)", type=['xlsx'], key=f"p1_hist_{st.session_state['uploader_key']}")

    with col2_p1:
        st.markdown("##### ⚙️ 2. 目標設定與演算法獨立指派")
        template_subjects_list = []
        selected_sheet_p1 = None
        if file_template_p1:
            try:
                temp_xls = pd.ExcelFile(file_template_p1)
                selected_sheet_p1 = st.selectbox("🎯 選擇要填入的出題總表工作表：", temp_xls.sheet_names, key="p1_sheet_select")
                df_temp_scan = pd.read_excel(file_template_p1, sheet_name=selected_sheet_p1, header=None).fillna("")
                unique_subjects = set()
                found_cols = []
                header_row = -1
                for r in range(0, min(15, len(df_temp_scan))):
                    current_row_cols = []
                    for c in range(len(df_temp_scan.columns)):
                        if str(df_temp_scan.iloc[r, c]).strip().replace(' ', '').replace(' ', '') == "科目": current_row_cols.append(c)
                    if len(current_row_cols) >= 3:
                        found_cols = current_row_cols; header_row = r; break
                if header_row != -1:
                    for c in found_cols:
                        for scan_r in range(header_row + 1, len(df_temp_scan)):
                            val = str(df_temp_scan.iloc[scan_r, c]).strip()
                            cleaned = clean_subject_name_p1(val)
                            if cleaned and cleaned not in ["科目", "編進度", "出題教師", "第一次", "第二次", "期末考", "補考", "教師", "共同科目"]:
                                if len(cleaned) <= 15 and not cleaned.startswith('↓') and not cleaned[0].isdigit(): unique_subjects.add(cleaned)
                template_subjects_list = sorted(list(unique_subjects))
            except: pass
                
        proportional_subjects = st.multiselect("👉 請選擇要套用【班級比例制 (教多出多)】的科目：", options=template_subjects_list, help="若有上傳歷史檔案，系統將自動扣除已出題次數。", key="p1_prop_select")
        if st.button("🗑️ 清除設定 (僅限階段一)", use_container_width=True, key="p1_clear"):
            st.session_state['results_p1'] = None
            st.session_state['debug_log_p1'] = []
            st.session_state['uploader_key'] += 1
            st.rerun()

    st.divider()

    if st.button("🚀 啟動出題教師智能排定", type="primary", use_container_width=True, key="btn_p1"):
        if not file_peike_p1 or not file_template_p1: st.error("🚨 請確認【配課表】與【出題總表範本】皆已上傳！")
        else:
            with st.spinner("🧠 啟動半衰期對稱演算法與歷史記憶扣除系統..."):
                try:
                    debug_msgs = []
                    history_map = extract_history(file_history_p1)
                    if history_map: debug_msgs.append("🕰️ 成功載入上學期歷史紀錄！已啟動跨學期公平扣除引擎。")
                    
                    sync_group_map = {}
                    if file_sync_p1:
                        df_sync = pd.read_csv(file_sync_p1) if file_sync_p1.name.endswith('.csv') else pd.read_excel(file_sync_p1, header=None)
                        sync_group_map = parse_sync_file(df_sync)
                    
                    xls_peike = pd.ExcelFile(file_peike_p1)
                    dict_g1 = build_class_teacher_dict(pd.read_excel(xls_peike, sheet_name='一年級') if '一年級' in xls_peike.sheet_names else pd.DataFrame())
                    dict_g2 = build_class_teacher_dict(pd.read_excel(xls_peike, sheet_name='二年級') if '二年級' in xls_peike.sheet_names else pd.DataFrame())
                    dict_g3 = build_class_teacher_dict(pd.read_excel(xls_peike, sheet_name='三年級') if '三年級' in xls_peike.sheet_names else pd.DataFrame())
                    
                    wb = openpyxl.load_workbook(file_template_p1)
                    ws = wb[selected_sheet_p1 if selected_sheet_p1 else wb.sheetnames[0]]
                    
                    header_row = -1
                    grade_subj_cols = []
                    for r in range(1, 15):
                        current_row_cols = [c for c in range(1, 25) if str(ws.cell(row=r, column=c).value).strip().replace(' ', '').replace(' ', '') == "科目"]
                        if len(current_row_cols) >= 3:
                            header_row = r; grade_subj_cols = current_row_cols[:3]; break
                    
                    if header_row == -1 or len(grade_subj_cols) < 3:
                        st.error("🚨 無法在範本中精準找到一、二、三年級的「科目」表頭，請確認範本格式。")
                        st.stop()
                        
                    grade_mapping = [{'grade': 3, 'dict': dict_g3, 'col': grade_subj_cols[0]}, {'grade': 2, 'dict': dict_g2, 'col': grade_subj_cols[1]}, {'grade': 1, 'dict': dict_g1, 'col': grade_subj_cols[2]}]
                    cell_tasks = []
                    group_teacher_pool = {} 
                    
                    for r in range(header_row + 1, ws.max_row + 1):
                        for mapping in grade_mapping:
                            subj_col = mapping['col']
                            cell_subj = ws.cell(row=r, column=subj_col)
                            subj_raw = cell_subj.value
                            if not subj_raw or str(subj_raw).strip() == "": continue
                                
                            subj_clean = clean_subject_name_p1(subj_raw)
                            grade_num = mapping['grade']
                            t_dict = mapping['dict']
                            
                            teachers = get_teachers_for_subject_p1(subj_raw, t_dict)
                            if subj_clean not in proportional_subjects: teachers = list(dict.fromkeys(teachers)) 
                            
                            group_id = sync_group_map.get((grade_num, subj_clean), None)
                            if group_id is not None:
                                if group_id not in group_teacher_pool: group_teacher_pool[group_id] = []
                                group_teacher_pool[group_id].extend(teachers)
                            
                            cell_tasks.append({
                                'row': r, 'subj_col': subj_col, 'mapping': mapping,
                                'group_id': group_id, 'teachers': teachers, 'subj_clean': subj_clean,
                                'raw_name': str(subj_raw).strip() 
                            })
                            
                    assignment_cache = {} 
                    cells_written_count = 0
                    global_teacher_assignment_counts = defaultdict(lambda: defaultdict(int))
                    
                    for task in cell_tasks:
                        final_teachers = task['teachers']
                        cache_key = task['group_id'] if task['group_id'] is not None else (task['mapping']['grade'], task['raw_name'])
                        subj_clean = task['subj_clean']
                        
                        if task['group_id'] is not None:
                            final_teachers = group_teacher_pool[task['group_id']]
                            
                        if subj_clean in proportional_subjects:
                            past_teachers = history_map.get((task['mapping']['grade'], task['raw_name']), [])
                            current_pool = final_teachers.copy()
                            for pt in past_teachers:
                                if pt in current_pool: current_pool.remove(pt)
                            final_teachers = current_pool if current_pool else final_teachers.copy()
                        else:
                            final_teachers = list(dict.fromkeys(final_teachers))
                        
                        if cache_key not in assignment_cache:
                            assigned_seq = generate_perfect_balanced_sequence(final_teachers, global_teacher_assignment_counts, sequence_length=10)
                            assignment_cache[cache_key] = assigned_seq
                        
                        assigned_teachers_sequence = assignment_cache[cache_key]
                            
                        if any(assigned_teachers_sequence):
                            r, subj_col, mapping = task['row'], task['subj_col'], task['mapping']
                            mode_str = "按勞(上下對稱)" if subj_clean in proportional_subjects else "平均"
                            
                            seq_idx = 0
                            for offset in range(1, 8):
                                target_col = subj_col + offset
                                if mapping != grade_mapping[-1] and target_col >= grade_mapping[grade_mapping.index(mapping)+1]['col']: break
                                if mapping == grade_mapping[-1] and target_col >= grade_subj_cols[-1] + 6: break
                                
                                cell_target = ws.cell(row=r, column=target_col)
                                val_check = cell_target.value
                                if val_check is None or str(val_check).strip() == "" or str(val_check).strip() == "None":
                                    if type(cell_target).__name__ != 'MergedCell':
                                        cell_target.value = str(assigned_teachers_sequence[seq_idx])
                                        cells_written_count += 1
                                seq_idx += 1
                                
                            if cells_written_count <= 25 and task['mapping']['grade'] == 1 and '英' in task['raw_name']:
                                debug_msgs.append(f"✍️ 寫入[{mode_str}]：{mapping['grade']}年級 {task['raw_name']} -> {assigned_teachers_sequence[:5]}")

                    out_bytes = io.BytesIO()
                    wb.save(out_bytes)
                    st.session_state['results_p1'] = out_bytes.getvalue()
                    st.session_state['debug_log_p1'] = debug_msgs
                    if cells_written_count > 0:
                        st.balloons()
                        st.success(f"🎉 大滿貫完成！填入了 {cells_written_count} 個欄位。跨科防撞與職普分流已完美發揮作用！")
                    else:
                        st.warning("⚠️ 系統已跑完運算，但沒有填入任何欄位。")
                except Exception as e:
                    st.error(f"🚨 發生錯誤：{e}")
                    st.code(traceback.format_exc())

    if st.session_state['results_p1']:
        st.divider()
        c_d1, c_d2 = st.columns([2, 1], gap="large")
        with c_d1:
            st.download_button("📥 下載全自動排定之【出題教師總表】", data=st.session_state['results_p1'], file_name="進度及出題教師總表_上下學期均分版.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, type="primary")
        with c_d2:
            with st.expander("🔎 系統寫入透視日誌 (點此展開)"):
                for msg in st.session_state.get('debug_log_p1', []): st.write(msg)


# ---------------------------------------------------------
# 【階段二：試卷催繳通知單】
# ---------------------------------------------------------
with tab2:
    st.subheader("📑 階段二：試卷催繳通知單自動生成系統")
    st.markdown("上傳包含催繳名單的 Excel，**選擇對應的工作表 (考試類型)**，系統會自動產出專屬的 Word 通知單。")
    
    if not HAS_DOCX: st.error("🚨 偵測到系統未安裝 `python-docx` 套件！請在環境中安裝 `python-docx`。")

    col1_p2, col2_p2 = st.columns([1, 1], gap="large")
    with col1_p2:
        st.markdown("##### ⚙️ 參數設定")
        deadline = st.text_input("📅 繳交截止日", value="6/26", help="例如：6/26", key="p2_deadline")
        sender_name = st.text_input("✍️ 發送人署名", value="試務組 耀中", key="p2_sender")
        
    with col2_p2:
        st.markdown("##### 📂 資料上傳與選擇")
        uploaded_file_p2 = st.file_uploader("請上傳「試卷催繳名單」(Excel)", type=["xlsx", "xls"], key=f"p2_uploader_{st.session_state['uploader_key']}")
        
        selected_sheet_p2 = None
        if uploaded_file_p2 is not None:
            try:
                excel_file_p2 = pd.ExcelFile(uploaded_file_p2)
                selected_sheet_p2 = st.selectbox("👇 請選擇工作表 (考試類型)：", excel_file_p2.sheet_names, key="p2_sheet_select")
            except: st.error("無法讀取 Excel 檔案。")

    if st.button("🚀 一鍵產出雙版本催繳通知單", use_container_width=True, type="primary", key="btn_p2"):
        if not HAS_DOCX: st.error("缺少 python-docx 套件。")
        elif not uploaded_file_p2 or not selected_sheet_p2: st.warning("⚠️ 請先上傳名單檔案，並選擇工作表！")
        else:
            try:
                df = pd.read_excel(uploaded_file_p2, sheet_name=selected_sheet_p2).dropna(how='all')
                if any(c not in df.columns for c in ['年級', '科目名稱', '姓名']):
                    st.error("🚨 上傳的檔案缺少必備欄位：年級、科目名稱、姓名。")
                else:
                    df['姓名'] = df['姓名'].astype(str).str.strip().replace('nan', '')
                    df['科目名稱'] = df['科目名稱'].astype(str).str.strip().replace('nan', '')
                    df['年級'] = df['年級'].astype(str).str.strip().replace('nan', '')
                    df = df[df['姓名'] != '']
                    
                    doc_print = Document()
                    doc_msg = Document()
                    
                    grouped = df.groupby('姓名')
                    for idx, (name, group) in enumerate(grouped):
                        exam_type = selected_sheet_p2
                        count = len(group)
                        
                        table = doc_print.add_table(rows=1, cols=1)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.style = 'Table Grid'
                        cell = table.cell(0, 0)
                        p_title_print = cell.paragraphs[0]
                        p_title_print.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_title_print = p_title_print.add_run(f"【{exam_type}】催繳試卷通知單")
                        run_title_print.bold = True
                        run_title_print.font.size = Pt(20) 
                        doc_print.add_paragraph()
                        
                        p_title_msg = doc_msg.add_paragraph()
                        p_title_msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_title_msg = p_title_msg.add_run(f"【{exam_type}】催繳試卷通知單")
                        run_title_msg.bold = True
                        run_title_msg.font.size = Pt(16)
                        doc_msg.add_paragraph()

                        for doc in [doc_print, doc_msg]:
                            doc.add_paragraph(f"{name} 老師您好:\\n")
                            doc.add_paragraph(f"{exam_type}試卷繳交截止日 {deadline} 已過，尚有 {count} 份試卷未繳:\\n")
                            for grade, grade_group in group.groupby('年級'):
                                doc.add_paragraph(f"[{grade}年級]")
                                for i, (_, row) in enumerate(grade_group.iterrows(), 1):
                                    doc.add_paragraph(f"  {i}. 科目: {row['科目名稱']}")
                            doc.add_paragraph(f"\\n{sender_name}")
                        
                        if idx < len(grouped) - 1:
                            doc_print.add_page_break()
                            doc_msg.add_paragraph("\\n" + "=" * 40 + "\\n")
                    
                    out_stream_print, out_stream_msg = io.BytesIO(), io.BytesIO()
                    doc_print.save(out_stream_print); doc_msg.save(out_stream_msg)
                    st.session_state['docx_data_p2_print'] = out_stream_print.getvalue()
                    st.session_state['docx_data_p2_msg'] = out_stream_msg.getvalue()
                    st.session_state['processed_p2'] = True
            except Exception as e:
                st.error(f"發生未預期錯誤: {e}"); st.code(traceback.format_exc())

    if st.session_state['processed_p2'] and st.session_state['docx_data_p2_print']:
        st.success(f"✅ 完美達成！已產出雙版本。")
        c1, c2 = st.columns([1, 1])
        with c1:
            st.download_button("🖨️ 下載：紙本列印版", st.session_state['docx_data_p2_print'], f"{selected_sheet_p2}催繳通知單_紙本版.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")
        with c2:
            st.download_button("💬 下載：訊息複製版", st.session_state['docx_data_p2_msg'], f"{selected_sheet_p2}催繳通知單_訊息版.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="secondary")


# ---------------------------------------------------------
# 【階段三：段考監考智能輔助系統】
# ---------------------------------------------------------
with tab3:
    st.subheader("📅 階段三：段考監考智能輔助系統 (終極完全體)")
    st.info("💡 終極升級：實裝「雙重智慧檢核機制」！AI 將在匯出總表全自動對帳，精準抓出人力缺口。")

    col1_p3, col2_p3 = st.columns([1, 1], gap="large")

    with col1_p3:
        st.markdown("##### 📂 1. 上傳排考與標籤資料")
        file_quota_p3 = st.file_uploader("1️⃣ 監考堂數.xlsx", type=['xlsx'], key=f"f1_p3_{st.session_state['uploader_key']}")
        file_list_p3 = st.file_uploader("2️⃣ 監考名單.xlsx", type=['xlsx'], key=f"f2_p3_{st.session_state['uploader_key']}")
        file_type_p3 = st.file_uploader("3️⃣ 監考類型總數.xlsx", type=['xlsx'], key=f"f3_p3_{st.session_state['uploader_key']}")
        file_pub_p3 = st.file_uploader("4️⃣ 監考總表公布版.xlsx (範本)", type=['xlsx'], key=f"f4_p3_{st.session_state['uploader_key']}")
        file_assign_p3 = st.file_uploader("5️⃣ 監考一覽表.xlsx (班級分配範本)", type=['xlsx'], key=f"f5_p3_{st.session_state['uploader_key']}")
        st.write("---")
        file_course_p3 = st.file_uploader("6️⃣ 配課表.xlsx (多工作表)", type=['xlsx'], key=f"f6_p3_{st.session_state['uploader_key']}")
        file_label_p3 = st.file_uploader("7️⃣ 標籤列印.xlsx (試卷袋範本)", type=['xlsx'], key=f"f7_p3_{st.session_state['uploader_key']}")

    with col2_p3:
        st.markdown("##### ⚙️ 2. 考試設定與特許名單")
        selected_sheet_p3 = None
        if file_quota_p3:
            xls = pd.ExcelFile(file_quota_p3)
            selected_sheet_p3 = st.selectbox("👇 選擇考試項目：", xls.sheet_names, key="p3_sheet_select")
        
        flex_names_p3 = []
        teacher_list_p3 = []
        if file_list_p3:
            temp_df = pd.read_excel(file_list_p3, header=None).dropna(how='all').fillna("")
            for c in range(5):
                try:
                    lst = temp_df.iloc[2:, c].astype(str).str.strip().tolist()
                    lst = [t for t in lst if t != "" and t != "nan" and not str(t).isdigit()]
                    if len(lst) > 10:
                        teacher_list_p3 = lst; break
                except: pass
            flex_names_p3 = st.multiselect("🛡️ 優先時數不大於名單：", options=teacher_list_p3, key="p3_flex_names")

        class_list_p3 = []
        if file_assign_p3:
            df_assign_temp = pd.read_excel(file_assign_p3, header=None).dropna(how='all').fillna("")
            raw_list = df_assign_temp.iloc[:, 0].astype(str).str.strip().tolist()
            class_names_raw = [x for x in raw_list if x and not any(bad in x for bad in ["班級", "日期", "節次", "星期", "一覽表", "總表", "華南", "期中考", "註"])]
            class_list_p3 = [normalize_cls(c) for c in class_names_raw]

        st.write("")
        c_d0, c_d1, c_d2 = st.columns(3)
        with c_d0:
            has_manual_p3 = st.checkbox("📌 包含手排日", value=True, key="p3_has_manual")
            if has_manual_p3: d0_date_p3 = st.date_input("📅 手排日期", datetime.now(), key="p3_d0_date")
            else: d0_date_p3 = None
                
        with c_d1: d1_date_p3 = st.date_input("📅 AI Day1", datetime.now(), key="p3_d1_date")
        with c_d2: d2_date_p3 = st.date_input("📅 AI Day2", datetime.now(), key="p3_d2_date")
        
        st.write("---")
        st.markdown("#### ⏳ 兼課教師可用時段精確鎖定")
        edited_time_df = st.data_editor(
            st.session_state.time_rules, 
            num_rows="dynamic", 
            use_container_width=True,
            column_config={
                "老師": st.column_config.SelectboxColumn("👨‍🏫 指定老師", options=teacher_list_p3 if teacher_list_p3 else [""], required=False),
                "允許日期": st.column_config.SelectboxColumn("📅 允許日期", options=["無限制", "僅 Day 1", "僅 Day 2"], default="無限制"),
                "允許節次": st.column_config.TextColumn("⏰ 允許節次 (如: 1,2,3)")
            },
            key="p3_time_editor"
        )

        st.write("---")
        st.markdown("#### 🎯 特定班級與老師綁定")

        file_bind_p3 = st.file_uploader("📥 [選填] 匯入既有綁定名單 (.xlsx)", type=['xlsx'], key=f"f_bind_{st.session_state['uploader_key']}")
        if file_bind_p3 and st.session_state.last_bind_file != file_bind_p3.name:
            try:
                df_bind_up = pd.read_excel(file_bind_p3).dropna(how='all')
                if "老師" in df_bind_up.columns and "班級" in df_bind_up.columns:
                    st.session_state.bind_rules = df_bind_up[["老師", "班級"]]
                    st.session_state.last_bind_file = file_bind_p3.name
                    st.rerun() 
            except: st.error("讀取失敗。")

        edited_bind_df = st.data_editor(
            st.session_state.bind_rules, 
            num_rows="dynamic", 
            use_container_width=True,
            column_config={
                "老師": st.column_config.SelectboxColumn("👨‍🏫 指定老師", options=teacher_list_p3 if teacher_list_p3 else [""], required=False),
                "班級": st.column_config.SelectboxColumn("🏫 指定班級", options=class_list_p3 if class_list_p3 else [""], required=False)
            },
            key="p3_bind_editor"
        )
        
        bind_output = io.BytesIO()
        with pd.ExcelWriter(bind_output, engine='xlsxwriter') as writer:
            edited_bind_df.to_excel(writer, index=False, header=True)
        st.download_button("💾 儲存綁定名單", bind_output.getvalue(), "特定老師班級綁定名單.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key="p3_dl_bind")
        
        if st.button("🗑️ 清除設定 (僅限階段三)", use_container_width=True, key="p3_clear"):
            st.session_state['results_p3'] = None
            st.session_state['uploader_key'] += 1
            if 'bind_rules' in st.session_state: del st.session_state['bind_rules']
            if 'time_rules' in st.session_state: del st.session_state['time_rules']
            if 'last_bind_file' in st.session_state: del st.session_state['last_bind_file']
            st.rerun()

    st.divider()

    if st.button("🚀 啟動終極全自動排班系統", type="primary", use_container_width=True, key="btn_p3"):
        if not all([file_quota_p3, file_list_p3, file_type_p3, file_assign_p3]):
            st.error("🚨 請至少確認【1, 2, 3, 5】號基礎檔案皆已上傳！")
        else:
            try:
                df_quota = pd.read_excel(file_quota_p3, sheet_name=selected_sheet_p3).dropna(how='all').fillna("")
                quota_dict = {str(df_quota.iloc[r, 0]).strip(): int(float(str(df_quota.iloc[r, 1]).strip())) for r in range(df_quota.shape[0]) if str(df_quota.iloc[r, 0]).strip() and str(df_quota.iloc[r, 1]).strip()}
                
                df_list_raw = pd.read_excel(file_list_p3, header=None).dropna(how='all').fillna("")
                header_row_idx = 1
                for r in range(min(5, df_list_raw.shape[0])):
                    if any(k in str(df_list_raw.iloc[r, 1]).strip() for k in ["教師", "姓名", "老師"]):
                        header_row_idx = r; break
                
                teacher_col_idx, quota_col_in_list = 1, 2
                period_cols = [c for c in range(3, df_list_raw.shape[1]) if extract_period_num(str(df_list_raw.iloc[header_row_idx, c])) != -1]
                
                if len(period_cols) < 1: st.error("🚨 無法辨識節次。"); st.stop()
                total_periods = len(period_cols)
                date_row_idx = header_row_idx - 1 if header_row_idx > 0 else 0
                date_row_s = df_list_raw.iloc[date_row_idx, :].replace(['nan', ''], np.nan).ffill().fillna("")
                
                manual_cols, ai_period_cols = [], []
                for c in period_cols:
                    if has_manual_p3 and d0_date_p3 and matches_date(str(date_row_s[c]), d0_date_p3): manual_cols.append(c)
                    else: ai_period_cols.append(c)
                    
                ai_periods = len(ai_period_cols)
                if ai_periods == 0: st.error("🚨 找不到 AI 處理的欄位"); st.stop()
                    
                ai_period_nums = [extract_period_num(str(df_list_raw.iloc[header_row_idx, c])) for c in ai_period_cols]
                day_starts = [0]
                for j in range(1, ai_periods):
                    if ai_period_nums[j] <= ai_period_nums[j-1]: day_starts.append(j)

                df_type = pd.read_excel(file_type_p3, header=None).dropna(how='all').fillna("")
                req_matrix = {'△': [0]*ai_periods, '※': [0]*ai_periods}
                for i in range(len(df_type)):
                    row_name = str(df_type.iloc[i, 0]).strip()
                    if row_name in ['△', '※']:
                        req_list = [int(float(v)) for c in range(1, df_type.shape[1]) if (v:=str(df_type.iloc[i, c]).strip())]
                        req_padded = (req_list + [0]*total_periods)[:total_periods]
                        req_matrix[row_name] = [req_padded[period_cols.index(c)] for c in ai_period_cols]

                ai_date_strs = [d1_date_p3.strftime('%m月%d日'), d2_date_p3.strftime('%m月%d日')]
                header_df = df_list_raw.iloc[0:header_row_idx+1].copy().astype(str).replace('nan', '')
                if has_manual_p3 and d0_date_p3:
                    for mc in manual_cols: header_df.iloc[date_row_idx, mc] = d0_date_p3.strftime('%m月%d日')
                for j in range(ai_periods): header_df.iloc[date_row_idx, ai_period_cols[j]] = get_ai_date_str(j, day_starts, ai_date_strs)
                
                df_list = df_list_raw.iloc[header_row_idx+1:].copy()
                teachers = [str(x).strip() for x in df_list.iloc[:, teacher_col_idx] if pd.notna(x) and str(x).strip() not in ["", "nan"]]

                time_constraints = {}
                for _, row in edited_time_df.iterrows():
                    t_name = str(row['老師']).strip()
                    if t_name and t_name != 'None':
                        p_limit_str = str(row['允許節次']).strip()
                        time_constraints[t_name] = {'day': str(row['允許日期']).strip(), 'periods': [int(p) for p in re.findall(r'\d+', p_limit_str)] if p_limit_str else []}

                with st.spinner(f"🧠 PuLP 運算中 ({len(teachers)} 位教師)..."):
                    prob = pulp.LpProblem("Scheduling", pulp.LpMinimize)
                    vX = {i: {j: pulp.LpVariable(f"X_{i}_{j}", cat='Binary') for j in range(ai_periods)} for i in range(len(teachers))}
                    vY = {i: {j: pulp.LpVariable(f"Y_{i}_{j}", cat='Binary') for j in range(ai_periods)} for i in range(len(teachers))}
                    
                    d1_idx = list(range(day_starts[0], day_starts[1])) if len(day_starts) > 1 else list(range(ai_periods))
                    d2_idx = list(range(day_starts[1], ai_periods)) if len(day_starts) > 1 else []

                    penalty = 0
                    for i, t in enumerate(teachers):
                        tgt = int(quota_dict.get(t, 0))
                        act = pulp.lpSum([vX[i][k] + vY[i][k]*2 for k in range(ai_periods)])
                        dfct_pos, dfct_neg = pulp.LpVariable(f"dfct_pos_{i}", 0), pulp.LpVariable(f"dfct_neg_{i}", 0)
                        prob += act + dfct_neg - dfct_pos == tgt
                        penalty += (dfct_pos + dfct_neg) * 500
                        if t in flex_names_p3: penalty -= dfct_neg * 400
                        
                        is_time_constrained = t in time_constraints
                        if is_time_constrained:
                            tc = time_constraints[t]
                            for j in range(ai_periods):
                                if (tc['day'] == '僅 Day 1' and j in d2_idx) or (tc['day'] == '僅 Day 2' and j in d1_idx) or (tc['periods'] and ai_period_nums[j] not in tc['periods']):
                                    prob += vX[i][j] == 0; prob += vY[i][j] == 0
                        
                        if tgt >= 5 and len(day_starts) >= 2 and not is_time_constrained:
                            prob += pulp.lpSum([vX[i][j] + vY[i][j] for j in d1_idx]) >= 1
                            prob += pulp.lpSum([vX[i][j] + vY[i][j] for j in d2_idx]) >= 1
                            prob += pulp.lpSum([vX[i][j] for j in range(ai_periods)]) >= 1
                            prob += pulp.lpSum([vY[i][j] for j in range(ai_periods)]) >= 1

                        for j in range(ai_periods):
                            prob += vX[i][j] + vY[i][j] <= 1
                            if (cv:=str(df_list.iloc[i, ai_period_cols[j]]).strip()) not in ["", "nan"]:
                                prob += vX[i][j] == 0; prob += vY[i][j] == 0
                            if ai_period_nums[j] in [3, 5]: prob += vX[i][j] == 0
                                
                        for j in range(ai_periods - 1):
                            if ai_period_nums[j] == 1 and ai_period_nums[j+1] == 2: prob += vX[i][j+1] >= vY[i][j]
                                
                    for j in range(ai_periods):
                        req_d, req_s = req_matrix['△'][j], req_matrix['※'][j]
                        slk_d_pos, slk_d_neg = pulp.LpVariable(f"slkd_pos_{j}", 0), pulp.LpVariable(f"slkd_neg_{j}", 0)
                        prob += pulp.lpSum([vX[i][j] for i in range(len(teachers))]) + slk_d_neg - slk_d_pos == req_d
                        penalty += (slk_d_pos + slk_d_neg) * 10000
                        slk_s_pos, slk_s_neg = pulp.LpVariable(f"slks_pos_{j}", 0), pulp.LpVariable(f"slks_neg_{j}", 0)
                        prob += pulp.lpSum([vY[i][j] for i in range(len(teachers))]) + slk_s_neg - slk_s_pos == req_s
                        penalty += (slk_s_pos + slk_s_neg) * 10000
                        
                    prob += penalty
                    prob.solve(pulp.PULP_CBC_CMD(timeLimit=45, msg=False, threads=1))

                    schedule_dict = {}
                    df_out_master = df_list.copy()
                    actual_matrix = {'△': [0]*ai_periods, '※': [0]*ai_periods}
                    
                    for i, t in enumerate(teachers):
                        res = []
                        df_out_master.iloc[i, quota_col_in_list] = int(quota_dict.get(t, 0))
                        for j in range(ai_periods):
                            val = str(df_list.iloc[i, ai_period_cols[j]]).strip()
                            if val in ["", "nan"]:
                                if vX[i][j].varValue == 1: val = "△"; actual_matrix['△'][j] += 1
                                elif vY[i][j].varValue == 1: val = "※"; actual_matrix['※'][j] += 1
                                else: val = "" 
                            else:
                                if val == "△": actual_matrix['△'][j] += 1
                                elif val == "※": actual_matrix['※'][j] += 1
                            res.append(val); df_out_master.iloc[i, ai_period_cols[j]] = val
                        schedule_dict[t] = res

                # 檢核區
                discrepancies = []
                empty_row = {c: "" for c in df_out_master.columns}
                row_act_d, row_req_d, row_act_s, row_req_s, row_diff = empty_row.copy(), empty_row.copy(), empty_row.copy(), empty_row.copy(), empty_row.copy()
                empty_row[df_out_master.columns[teacher_col_idx]] = "--- 系統自動檢核區 ---"
                row_act_d[df_out_master.columns[teacher_col_idx]] = "實際排入 (△)"
                row_req_d[df_out_master.columns[teacher_col_idx]] = "需求總數 (△)"
                row_act_s[df_out_master.columns[teacher_col_idx]] = "實際排入 (※)"
                row_req_s[df_out_master.columns[teacher_col_idx]] = "需求總數 (※)"
                row_diff[df_out_master.columns[teacher_col_idx]]  = "異常差額警示"

                for j in range(ai_periods):
                    col_name = df_out_master.columns[ai_period_cols[j]]
                    period_name = str(df_list_raw.iloc[header_row_idx, ai_period_cols[j]]).strip()
                    day_name = "Day1" if j in d1_idx else "Day2"
                    act_d, req_d = actual_matrix['△'][j], req_matrix['△'][j]
                    act_s, req_s = actual_matrix['※'][j], req_matrix['※'][j]
                    row_act_d[col_name], row_req_d[col_name] = act_d, req_d
                    row_act_s[col_name], row_req_s[col_name] = act_s, req_s
                    diff_d, diff_s = act_d - req_d, act_s - req_s
                    diff_strs = []
                    if diff_d != 0: diff_strs.append(f"△{'+' if diff_d>0 else ''}{diff_d}"); discrepancies.append(f"【{day_name}】第 {period_name} 節 - △: 需求 {req_d} 人, 實際 {act_d} 人 (差額: {diff_d})")
                    if diff_s != 0: diff_strs.append(f"※{'+' if diff_s>0 else ''}{diff_s}"); discrepancies.append(f"【{day_name}】第 {period_name} 節 - ※: 需求 {req_s} 人, 實際 {act_s} 人 (差額: {diff_s})")
                    row_diff[col_name] = "、".join(diff_strs) if diff_strs else "正常吻合"

                df_out_master = pd.concat([pd.DataFrame([empty_row, row_act_d, row_req_d, row_act_s, row_req_s, row_diff, {c: "" for c in df_out_master.columns}]), df_out_master], ignore_index=True)

                with st.spinner("🎯 執行班級自動分配..."):
                    df_assign_calc = pd.read_excel(file_assign_p3, header=None).dropna(how='all').fillna("")
                    class_names_raw = [x for x in df_assign_calc.iloc[:, 0].astype(str).str.strip().tolist() if x and not any(bad in x for bad in ["班級", "日期", "節次", "星期", "一覽表", "總表", "華南", "期中考", "註"])]
                    assign_map = {normalize_cls(name): idx for idx, name in enumerate(class_names_raw)}
                    
                    t2c_map = {str(row['老師']).strip(): assign_map[normalize_cls(row['班級'])] for _, row in edited_bind_df.iterrows() if str(row['老師']).strip() and str(row['老師']).strip() != 'None' and normalize_cls(row['班級']) in assign_map}
                    assigned_matrix = np.empty((len(class_names_raw), ai_periods), dtype=object)
                    
                    for i_day, day_start in enumerate(day_starts):
                        day_end = day_starts[i_day+1] if i_day+1 < len(day_starts) else ai_periods
                        day_length = day_end - day_start
                        j1 = day_start
                        proctors_j1 = [t for t in teachers if schedule_dict[t][j1] in ["△", "※"]]
                        random.shuffle(proctors_j1)
                        rem_j1 = []
                        for p in proctors_j1:
                            if p in t2c_map and assigned_matrix[t2c_map[p], j1] is None: assigned_matrix[t2c_map[p], j1] = p
                            else: rem_j1.append(p)
                        r_ptr = 0
                        for idx in range(len(class_names_raw)):
                            if assigned_matrix[idx, j1] is None and r_ptr < len(rem_j1): assigned_matrix[idx, j1] = rem_j1[r_ptr]; r_ptr += 1
                        
                        if day_length > 1:
                            j2 = day_start + 1
                            bound = {p_prev: True for idx in range(len(class_names_raw)) if (p_prev:=assigned_matrix[idx, j1]) in schedule_dict and schedule_dict[p_prev][j1] == "※" and schedule_dict[p_prev][j2] == "△"}
                            rem = [p for p in [t for t in teachers if schedule_dict[t][j2] in ["△", "※"]] if p not in bound]
                            random.shuffle(rem)
                            rem_after_bind = []
                            for p in rem:
                                if p in t2c_map and assigned_matrix[t2c_map[p], j2] is None: assigned_matrix[t2c_map[p], j2] = p
                                else: rem_after_bind.append(p)
                            for p_prev in bound.keys():
                                for idx in range(len(class_names_raw)):
                                    if assigned_matrix[idx, j1] == p_prev: assigned_matrix[idx, j2] = p_prev
                            r_idx = 0
                            for idx in range(len(class_names_raw)):
                                if assigned_matrix[idx, j2] is None and r_idx < len(rem_after_bind): assigned_matrix[idx, j2] = rem_after_bind[r_idx]; r_idx += 1

                            for offset in range(2, day_length):
                                curr_j = day_start + offset
                                proctors = [t for t in teachers if schedule_dict[t][curr_j] in ["△", "※"]]
                                random.shuffle(proctors)
                                rem_curr = []
                                for p in proctors:
                                    if p in t2c_map and assigned_matrix[t2c_map[p], curr_j] is None: assigned_matrix[t2c_map[p], curr_j] = p
                                    else: rem_curr.append(p)
                                r_ptr = 0
                                for idx in range(len(class_names_raw)):
                                    if assigned_matrix[idx, curr_j] is None and r_ptr < len(rem_curr): assigned_matrix[idx, curr_j] = rem_curr[r_ptr]; r_ptr += 1

                    class_proctor_schedule = {normalize_cls(c_name): [assigned_matrix[r_idx, col] for col in range(ai_periods)] for r_idx, c_name in enumerate(class_names_raw)}

                    wb_assign = openpyxl.load_workbook(file_assign_p3)
                    ws_assign = wb_assign.active
                    manual_proctors = {} 
                    first_class_row, class_col_idx = -1, 1
                    for r in range(1, 20):
                        for c in range(1, 5):
                            if (v:=ws_assign.cell(row=r, column=c).value) and str(v).strip() in class_names_raw: first_class_row, class_col_idx = r, c; break
                        if first_class_row != -1: break
                    
                    if first_class_row != -1:
                        target_cols = [class_col_idx + 1 + i for i in range(total_periods)]
                        manual_assign_cols, ai_assign_cols = [], []
                        for i, c in enumerate(period_cols):
                            if c in manual_cols: manual_assign_cols.append(target_cols[i])
                            else: ai_assign_cols.append(target_cols[i])
                                
                        date_row = next((r for r in range(first_class_row - 1, max(0, first_class_row - 4), -1) if (val:=str(ws_assign.cell(row=r, column=target_cols[0]).value).strip()) != "" and not val.isdigit() and "期中" not in val and "華南" not in val), -1)
                        if date_row != -1:
                            if has_manual_p3 and d0_date_p3:
                                for mc_a in manual_assign_cols:
                                    try: ws_assign.cell(row=date_row, column=mc_a).value = d0_date_p3.strftime('%m月%d日')
                                    except: pass
                            for j in range(ai_periods):
                                try: ws_assign.cell(row=date_row, column=ai_assign_cols[j]).value = get_ai_date_str(j, day_starts, ai_date_strs)
                                except: pass
                        
                        for r in range(first_class_row, ws_assign.max_row + 1):
                            if c_val:=ws_assign.cell(row=r, column=class_col_idx).value:
                                norm_c = normalize_cls(c_val)
                                if has_manual_p3:
                                    if norm_c not in manual_proctors: manual_proctors[norm_c] = {}
                                    for mc_idx, mc_a in enumerate(manual_assign_cols):
                                        p_val = extract_period_num(str(df_list_raw.iloc[header_row_idx, manual_cols[mc_idx]]))
                                        if val_m:=ws_assign.cell(row=r, column=mc_a).value: manual_proctors[norm_c][p_val] = str(val_m).strip()
                                if norm_c in class_proctor_schedule:
                                    for j in range(ai_periods): ws_assign.cell(row=r, column=ai_assign_cols[j]).value = class_proctor_schedule[norm_c][j]
                    
                    out_assign = io.BytesIO()
                    wb_assign.save(out_assign)
                    assign_bytes = out_assign.getvalue()

                pub_bytes = None
                if file_pub_p3:
                    with st.spinner("🖨️ 正在無縫套印至公布版..."):
                        wb = openpyxl.load_workbook(file_pub_p3)
                        ws = wb.active
                        h_row = next((r for r in range(1, 16) for c in range(1, 61) if (val:=ws.cell(row=r, column=c).value) and any(k in str(val) for k in ["教師", "姓名", "老師"])), -1)
                        if h_row != -1:
                            t_cols = [c for c in range(1, 61) if (val:=ws.cell(row=h_row, column=c).value) and any(k in str(val) for k in ["教師", "姓名", "老師"])]
                            for c in t_cols:
                                t_col_target = [scan_c for scan_c in range(c + 1, c + 25) if (val:=str(ws.cell(row=h_row, column=scan_c).value).strip()) and not any(k in val for k in ["教師", "姓名", "標號", "老師"]) and extract_period_num(val) != -1]
                                if len(t_col_target) >= total_periods:
                                    pub_manual_cols, pub_ai_cols = [], []
                                    for i, pc in enumerate(period_cols):
                                        if pc in manual_cols: pub_manual_cols.append(t_col_target[i])
                                        else: pub_ai_cols.append(t_col_target[i])
                                    if has_manual_p3 and d0_date_p3:
                                        for pmc in pub_manual_cols:
                                            try: ws.cell(row=h_row-1, column=pmc).value = d0_date_p3.strftime('%m月%d日')
                                            except: pass
                                    for j in range(ai_periods):
                                        try: ws.cell(row=h_row-1, column=pub_ai_cols[j]).value = get_ai_date_str(j, day_starts, ai_date_strs)
                                        except: pass
                                    for r in range(h_row+1, ws.max_row + 1):
                                        if t_val:=ws.cell(row=r, column=c).value:
                                            if (name:=str(t_val).strip()) in schedule_dict:
                                                for j in range(ai_periods): ws.cell(row=r, column=pub_ai_cols[j]).value = schedule_dict[name][j]
                        out_pub = io.BytesIO()
                        wb.save(out_pub)
                        pub_bytes = out_pub.getvalue()

                label_bytes = None
                if file_course_p3 and file_label_p3:
                    with st.spinner("🏷️ 正在合成試卷袋標籤..."):
                        course_dict = {}
                        xls_course = pd.ExcelFile(file_course_p3)
                        for sheet in xls_course.sheet_names:
                            df_c = pd.read_excel(file_course_p3, sheet_name=sheet).dropna(how='all').fillna("")
                            for r_idx, row in df_c.iterrows():
                                subj_raw = str(row.iloc[0]).strip()
                                if not subj_raw: continue
                                subj_norm = normalize_subject_p3(subj_raw)
                                for c_idx in range(1, len(df_c.columns)):
                                    cls_raw = str(df_c.columns[c_idx]).strip()
                                    teacher = str(row.iloc[c_idx]).strip()
                                    if teacher and cls_raw: course_dict[(normalize_cls(cls_raw), subj_norm)] = teacher
                        
                        wb_label = openpyxl.load_workbook(file_label_p3)
                        ws_label = wb_label.active
                        col_map = {}
                        header_row = 1
                        for r in range(1, 6):
                            for c in range(1, ws_label.max_column + 1):
                                val = str(ws_label.cell(row=r, column=c).value).strip()
                                if "班級" in val and '班級' not in col_map: col_map['班級'] = c
                                elif "科目" in val and '科目' not in col_map: col_map['科目'] = c
                                elif "日期" in val and '日期' not in col_map: col_map['日期'] = c
                                elif "序號" in val and '序號' not in col_map: col_map['序號'] = c
                                elif "任課" in val and '任課教師' not in col_map: col_map['任課教師'] = c
                                elif "監考" in val and '監考老師' not in col_map: col_map['監考老師'] = c
                            if '班級' in col_map and '監考老師' in col_map: header_row = r; break

                        d1_ymd, d1_short, d1_slash = d1_date_p3.strftime('%Y-%m-%d'), d1_date_p3.strftime('%m-%d'), d1_date_p3.strftime('%Y/%m/%d')
                        d2_ymd, d2_short, d2_slash = d2_date_p3.strftime('%Y-%m-%d'), d2_date_p3.strftime('%m-%d'), d2_date_p3.strftime('%Y/%m/%d')
                        if has_manual_p3 and d0_date_p3: d0_ymd, d0_short, d0_slash = d0_date_p3.strftime('%Y-%m-%d'), d0_date_p3.strftime('%m-%d'), d0_date_p3.strftime('%Y/%m/%d')

                        day_p_val_to_ai_col = {}
                        curr_day_idx = 0
                        for j in range(ai_periods):
                            if j in day_starts and j != 0: curr_day_idx += 1
                            day_p_val_to_ai_col[(curr_day_idx, ai_period_nums[j])] = j

                        for r in range(header_row + 1, ws_label.max_row + 1):
                            if '班級' not in col_map: continue
                            cls_raw = ws_label.cell(row=r, column=col_map['班級']).value
                            if cls_raw is None or not str(cls_raw).strip(): continue
                            
                            subj_raw = ws_label.cell(row=r, column=col_map['科目']).value if '科目' in col_map else ""
                            date_val = ws_label.cell(row=r, column=col_map['日期']).value if '日期' in col_map else ""
                            date_str = date_val.strftime('%Y-%m-%d') if isinstance(date_val, datetime) else str(date_val).split()[0].strip() if date_val is not None else ""
                            seq_val = ws_label.cell(row=r, column=col_map['序號']).value if '序號' in col_map else ""
                            
                            cls = normalize_cls(cls_raw)
                            subj = normalize_subject_p3(subj_raw)
                            
                            if '任課教師' in col_map:
                                teacher = course_dict.get((cls, subj), "")
                                if not teacher:
                                    for (c, s), t in course_dict.items():
                                        if c == cls and (subj in s or s in subj): teacher = t; break
                                if teacher: ws_label.cell(row=r, column=col_map['任課教師']).value = teacher
                            
                            try: p_val = int(float(str(seq_val).strip()))
                            except: p_val = -1
                            
                            if '監考老師' in col_map:
                                if has_manual_p3 and d0_date_p3 and any(d in date_str for d in [d0_ymd, d0_short, d0_slash]):
                                    if cls in manual_day0_proctors: ws_label.cell(row=r, column=col_map['監考老師']).value = manual_day0_proctors[cls]
                                elif cls in class_proctor_schedule and p_val != -1:
                                    day_idx = -1
                                    if any(d in date_str for d in [d1_ymd, d1_short, d1_slash]): day_idx = 0
                                    elif any(d in date_str for d in [d2_ymd, d2_short, d2_slash]): day_idx = 1
                                    
                                    if day_idx != -1 and (day_idx, p_val) in day_p_val_to_ai_col:
                                        target_col = day_p_val_to_ai_col[(day_idx, p_val)]
                                        ws_label.cell(row=r, column=col_map['監考老師']).value = class_proctor_schedule[cls][target_col]

                        out_label = io.BytesIO()
                        wb_label.save(out_label)
                        label_bytes = out_label.getvalue()

                st.session_state['results_p3'] = {
                    'orig': to_excel_bytes(df_out_master, header_df),
                    'assign': assign_bytes,
                    'pub': pub_bytes,
                    'label': label_bytes,
                    'discrepancies': discrepancies
                }
                
                if not discrepancies:
                    st.balloons()
                    st.success("✅ 完美排班！所有節次的監考人數與「監考類型總數」100% 吻合！")
                else:
                    st.warning("⚠️ 檢核提示：因特定鎖定條件，部分節次排入人數與需求有落差，明細如下：")
                    for d in discrepancies: st.write(f"- {d}")
                    st.info("💡 匯出的總表最上方有完整對帳明細。")

            except Exception as e:
                st.error("🚨 **系統發生錯誤！**"); st.code(traceback.format_exc(), language="python")

    if st.session_state['results_p3']:
        st.divider()
        res = st.session_state['results_p3']
        c1, c2, c3, c4 = st.columns(4)
        with c1: st.download_button("📥 1. 監考總表", res['orig'], "監考總表.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key="dl_p3_1")
        with c2: st.download_button("📥 2. 監考一覽表", res['assign'], "監考一覽表_分配完成.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, type="primary", key="dl_p3_2")
        with c3: 
            if res['pub']: st.download_button("📥 3. 公布版套印總表", res['pub'], "公布版總表.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key="dl_p3_3")
        with c4:
            if res.get('label'): st.download_button("📥 4. 標籤列印(完美接合)", res['label'], "標籤列印_完整版.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, type="primary", key="dl_p3_4")
