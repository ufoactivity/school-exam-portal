import streamlit as st
import pandas as pd
import numpy as np
import io
import re
import traceback
import datetime

# 嘗試載入 Word 排版套件，若無安裝則使用 Excel 完美分頁備案
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ==========================================
# 1. 網頁頁面配置
# ==========================================
st.set_page_config(page_title="補考自動化神器-頂規網頁版", page_icon="🏫", layout="wide")

st.title("📝 試務組-補考作業智能輔助系統")
st.info("💡 修正說明：新增單日/雙日補考勾選切換功能！學生名冊表頭日期將根據您的勾選自動智慧套印。")

if not HAS_DOCX:
    st.warning("💡 溫馨提醒：系統偵測未安裝 `python-docx`，已自動為您產出「Excel 列印分頁版」公告。若未來需要產出更精美的 Word 版，請在系統終端機輸入 `pip install python-docx` 後重啟網頁即可。")

# --- 初始化快取記憶體與清空鑰匙 ---
if 'results' not in st.session_state:
    st.session_state['results'] = None
if 'uploader_key' not in st.session_state:
    st.session_state['uploader_key'] = 0

# ==========================================
# 2. 輔助功能定義
# ==========================================
def get_str_col(df, keywords):
    if isinstance(keywords, str): keywords = [keywords]
    for kw in keywords:
        for i, col in enumerate(df.columns):
            if kw == str(col).strip():
                return df.iloc[:, i].fillna("").astype(str).str.strip()
    for kw in keywords:
        for i, col in enumerate(df.columns):
            if kw in str(col):
                return df.iloc[:, i].fillna("").astype(str).str.strip()
    return pd.Series([""] * len(df), index=df.index)

def grade_to_chinese(text):
    t = str(text)
    if '一' in t: return '一'
    if '二' in t: return '二'
    if '三' in t: return '三'
    if '1' in t or '１' in t: return '一'
    if '2' in t or '２' in t: return '二'
    if '3' in t or '３' in t: return '三'
    return "未知"

def natural_sort_key(s):
    return tuple(int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', str(s)))

def to_excel_bytes(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False)
    return output.getvalue()

# --- 姓名遮蔽處理函數 (個資保護) ---
def mask_name_func(name):
    n = str(name).strip()
    if len(n) <= 1: return n
    if len(n) == 2: return n[0] + "〇"
    return n[0] + "〇" + n[2:]

# --- 日期轉換函數 (轉換為 民國年與星期) ---
def format_tw_date(d):
    tw_year = d.year - 1911
    weekdays = ["一", "二", "三", "四", "五", "六", "日"]
    weekday_str = weekdays[d.weekday()]
    return f"{tw_year}年{d.month:02d}月{d.day:02d}日(星期{weekday_str})"

# --- 新增：Word 獨立分頁排版引擎 ---
def to_word_scope_bytes(df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微軟正黑體'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    classes = df['班級'].unique()
    for i, cls in enumerate(classes):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"【{cls}】補考科目及範圍公告")
        run.font.size = Pt(18)
        run.font.bold = True
        doc.add_paragraph()

        sub_df = df[df['班級'] == cls]
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '補考科目'
        hdr_cells[1].text = '補考範圍'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    
        for _, row in sub_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['所有補考的科目'])
            row_cells[1].text = str(row['補考範圍'])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        if i < len(classes) - 1:
            doc.add_page_break()
            
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# --- 新增：Excel 完美分頁排版引擎 (防呆備案) ---
def to_excel_scope_bytes(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        worksheet = workbook.add_worksheet('列印公告用')
        
        title_fmt = workbook.add_format({'bold': True, 'font_size': 18, 'align': 'center', 'valign': 'vcenter'})
        hdr_fmt = workbook.add_format({'bold': True, 'border': 1, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#D9D9D9', 'font_size': 12})
        subj_fmt = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter', 'font_size': 12})
        cell_fmt = workbook.add_format({'border': 1, 'align': 'left', 'valign': 'vcenter', 'text_wrap': True, 'font_size': 12})
        
        worksheet.set_column('A:A', 20)
        worksheet.set_column('B:B', 70)
        worksheet.set_margins(left=0.5, right=0.5, top=0.8, bottom=0.8)
        worksheet.fit_to_pages(1, 0)
        
        current_row = 0
        classes = df['班級'].unique()
        h_breaks = []
        
        for i, cls in enumerate(classes):
            sub_df = df[df['班級'] == cls]
            worksheet.merge_range(current_row, 0, current_row, 1, f"【{cls}】補考科目及範圍公告", title_fmt)
            worksheet.set_row(current_row, 30)
            current_row += 1
            
            worksheet.write(current_row, 0, "補考科目", hdr_fmt)
            worksheet.write(current_row, 1, "補考範圍", hdr_fmt)
            worksheet.set_row(current_row, 20)
            current_row += 1
            
            for _, row in sub_df.iterrows():
                worksheet.write(current_row, 0, str(row['所有補考的科目']), subj_fmt)
                worksheet.write(current_row, 1, str(row['補考範圍']), cell_fmt)
                text_len = len(str(row['補考範圍']))
                lines = str(row['補考範圍']).count('\n') + (text_len // 40) + 1
                worksheet.set_row(current_row, max(25, lines * 18))
                current_row += 1
                
            current_row += 1 
            if i < len(classes) - 1:
                h_breaks.append(current_row)
                
        worksheet.set_h_pagebreaks(h_breaks)
        
    return output.getvalue()

# --- ⭐ 新增：報表六 (學生名冊) 專屬精美排版匯出引擎 (支援單日判斷) ---
def to_excel_student_list_bytes(df, school_year, date1, date2=None):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        worksheet = workbook.add_worksheet('補考名冊公告')
        
        # 設定格式
        title_fmt = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter'})
        info_fmt = workbook.add_format({'font_size': 12, 'align': 'left', 'valign': 'vcenter'})
        hdr_fmt = workbook.add_format({'bold': True, 'border': 1, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#D9D9D9', 'font_size': 12})
        cell_fmt = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter', 'font_size': 12})
        
        # 智慧判斷單日或雙日，並轉換為民國年格式
        date_str1 = format_tw_date(date1)
        if date2 is not None:
            date_str2 = format_tw_date(date2)
            date_display_text = f"補考日期：{date_str1}、{date_str2}"
        else:
            date_display_text = f"補考日期：{date_str1}"
        
        # 寫入標題
        worksheet.merge_range(0, 0, 0, 6, f"{school_year}學年 全校補考名冊", title_fmt)
        worksheet.set_row(0, 25)
        
        # 寫入時間與地點資訊
        worksheet.write(2, 0, date_display_text, info_fmt)
        worksheet.write(3, 0, "補考時間：上午場 統一 八點十分開始；下午場 統一 一點十分開始", info_fmt)
        worksheet.write(4, 0, "補考地點：致用樓四樓會議室、圖書館三樓自修教室", info_fmt)
        
        # 寫入注意事項
        notes = [
            "注意事項：",
            "1.請同學穿著全套校服、攜帶學生證(或身分證)，準時應考。未符合前述規定者，不得補考。",
            "2.遲到15分鐘者不得入場，開始30分鐘後始能繳卷離場。",
            "3.入考場時，僅能攜帶證件、考試用文具，不得攜帶手機或任何電子設備。",
            "4.若有任何違規情事發生，將依本校考試規則辦理。",
            "5.如有問題請電洽教務處 試務組(分機164)。"
        ]
        
        row_idx = 6
        for note in notes:
            worksheet.write(row_idx, 0, note, info_fmt)
            row_idx += 1
            
        row_idx += 1  # 留一行空白
        start_row = row_idx
        
        # 寫入表格標題
        for col_num, col_name in enumerate(df.columns):
            worksheet.write(start_row, col_num, col_name, hdr_fmt)
            
        # 寫入表格內容
        for r_idx, r_data in enumerate(df.values):
            for c_idx, val in enumerate(r_data):
                worksheet.write(start_row + 1 + r_idx, c_idx, str(val), cell_fmt)
                
        # 調整列印版面與欄寬，讓報表美觀
        worksheet.set_column(0, 0, 12) # 班級
        worksheet.set_column(1, 1, 8)  # 座號
        worksheet.set_column(2, 2, 12) # 學號
        worksheet.set_column(3, 3, 12) # 姓名
        worksheet.set_column(4, 4, 30) # 科目
        worksheet.set_column(5, 5, 30) # 補考時間
        worksheet.set_column(6, 6, 20) # 補考地點
        
        worksheet.set_margins(left=0.5, right=0.5, top=0.8, bottom=0.8)
        worksheet.fit_to_pages(1, 0)
        
    return output.getvalue()

# ==========================================
# 3. 介面佈局：功能選單
# ==========================================
st.divider()
col_files, col_opts = st.columns([1, 1], gap="large")

with col_files:
    st.subheader("📂 第一步：上傳原始資料")
    file_target = st.file_uploader("1️⃣ 補考名單.xlsx", type=['xlsx'], key=f"f1_{st.session_state['uploader_key']}")
    file_short = st.file_uploader("2️⃣ 科目簡稱.xlsx", type=['xlsx'], key=f"f2_{st.session_state['uploader_key']}")
    file_exam = st.file_uploader("3️⃣ 科目對照表.xlsx", type=['xlsx'], key=f"f3_{st.session_state['uploader_key']}")
    file_teacher = st.file_uploader("4️⃣ 監考教師及時間.xlsx", type=['xlsx'], key=f"f4_{st.session_state['uploader_key']}")

with col_opts:
    st.subheader("⚙️ 第二步：考場容量與分流設定")
    zhiyong_cap = st.radio("📍 致用樓4樓會議室 人數上限：", [136, 148], index=0, horizontal=True)
    st.write("") 
    separate_mode = st.toggle("🔥 開啟【多科與單科嚴格分流】功能", value=False)
    
    st.divider()
    st.subheader("🗓️ 第三步：報表參數設定 (名冊公告用)")
    school_year_input = st.text_input("🎓 學年度 (例如：112-2)：", "112-2")
    
    # ⭐ 新增：單雙日切換功能
    is_two_days = st.checkbox("🗓️ 本次補考包含「第二天」", value=True)
    
    d_col1, d_col2 = st.columns(2)
    with d_col1:
        exam_date_1 = st.date_input("📅 補考日期(第一天)：", datetime.date.today())
    with d_col2:
        if is_two_days:
            exam_date_2 = st.date_input("📅 補考日期(第二天)：", datetime.date.today() + datetime.timedelta(days=1))
        else:
            exam_date_2 = None
        
    self_exam_deadline = st.date_input("📅 自行補考期限設定 (無試卷科目適用)：", datetime.date.today() + datetime.timedelta(days=7))

    st.write("")
    if st.button("🗑️ 清除舊資料 / 重新設定", use_container_width=True):
        st.session_state['results'] = None
        st.session_state['uploader_key'] += 1
        st.rerun()

# ==========================================
# 4. 執行與運算
# ==========================================
st.divider()

if st.button("🚀 開始智慧排考運算", type="primary", use_container_width=True):
    if not all([file_target, file_short, file_exam, file_teacher]):
        st.error("🚨 錯誤：請確認上方【4個檔案】皆已上傳完畢！")
    else:
        with st.spinner("系統正在執行智慧運算中..."):
            try:
                # 讀取 Excel
                df_short_map = pd.read_excel(file_short)
                df_exam_map = pd.read_excel(file_exam)
                df_target = pd.read_excel(file_target)
                df_teacher = pd.read_excel(file_teacher)

                grade_weight = {'一': 1, '二': 2, '三': 3}
                loc_weight = {'致用樓四樓會議室': 1, '圖書館三樓自修教室': 2, '電腦教室401': 3}
                current_targets = {"致用樓四樓會議室": zhiyong_cap, "圖書館三樓自修教室": 98, "電腦教室401": 37}

                # --- 階段一：基本資料處理 ---
                df_target['姓名'] = get_str_col(df_target, ['姓名', '學生姓名'])
                col_opencourse = get_str_col(df_target, ['開課班'])
                col_homeroom = get_str_col(df_target, ['班級', '原班級'])
                if col_opencourse.eq("").all(): col_opencourse = col_homeroom

                df_target['學號'] = get_str_col(df_target, ['學號'])
                df_target['座號'] = get_str_col(df_target, ['座號'])
                df_target['科目'] = get_str_col(df_target, ['科目', '考科'])
                df_target['班級'] = col_homeroom
                df_target['年級'] = df_target['班級'].apply(grade_to_chinese)

                col_a_s, col_b_s = df_short_map.columns[0], df_short_map.columns[1]
                short_dict = dict(zip(df_short_map[col_a_s].astype(str).str.strip(), df_short_map[col_b_s].astype(str).str.strip()))
                df_target['科目簡稱'] = df_target['科目'].map(short_dict).fillna("")
                
                ex_cls = get_str_col(df_exam_map, ['班級', '開課班'])
                ex_sub = get_str_col(df_exam_map, ['科目', '考科'])
                ex_pap = get_str_col(df_exam_map, ['試卷編號', '代碼'])
                if ex_pap.eq("").all() and df_exam_map.shape[1] > 7: ex_pap = df_exam_map.iloc[:, 7].astype(str).str.strip()
                ex_dict = dict(zip(ex_cls + ex_sub, ex_pap))
                
                df_target['試卷編號'] = (col_opencourse + df_target['科目']).map(ex_dict).fillna((col_homeroom + df_target['科目']).map(ex_dict)).fillna("")
                df_target['試卷編號'] = df_target['試卷編號'].apply(lambda x: str(x).replace('.0','') if str(x).endswith('.0') else str(x))

                # 將「補考範圍」存入 df_target 以供後續報表五使用
                ex_scope = get_str_col(df_exam_map, ['補考範圍', '範圍', '測驗範圍', '考試範圍'])
                scope_dict = dict(zip(ex_cls + ex_sub, ex_scope))
                df_target['補考範圍'] = (col_opencourse + df_target['科目']).map(scope_dict).fillna((col_homeroom + df_target['科目']).map(scope_dict)).fillna("")

                # 場地分配
                df_temp = df_target.drop_duplicates(subset=['學號', '試卷編號'], keep='first')
                v_counts = df_temp[df_temp['試卷編號'] != ""].groupby('學號').size()
                df_target['科目數目'] = df_target['學號'].map(v_counts).fillna(0).astype(int)
                
                df_students = df_target.drop_duplicates(subset=['學號']).copy()
                df_students = df_students[df_students['科目數目'] > 0].sort_values(by=['年級', '科目數目', '班級', '座號'], ascending=[True, False, True, True])

                venue_map = {}
                for gr, group in df_students.groupby('年級'):
                    if separate_mode:
                        multi = group[group['科目數目'] >= 2]
                        single = group[group['科目數目'] == 1]
                        m_v = (['致用樓四樓會議室'] * zhiyong_cap + ['圖書館三樓自修教室'] * 98 + ['電腦教室401'] * 37)
                        m_ans = m_v[:len(multi)]
                        rem_l = max(0, 98 - m_ans.count('圖書館三樓自修教室'))
                        rem_c = max(0, 37 - m_ans.count('電腦教室401'))
                        s_v = (['圖書館三樓自修教室'] * rem_l + ['電腦教室401'] * rem_c)
                        s_ans = s_v[:len(single)]
                        for sid, v in zip(multi['學號'], m_ans): venue_map[sid] = v
                        for sid, v in zip(single['學號'], s_ans): venue_map[sid] = v
                    else:
                        vns = (['致用樓四樓會議室'] * zhiyong_cap + ['圖書館三樓自修教室'] * 98 + ['電腦教室401'] * 37)
                        for sid, v in zip(group['學號'], vns[:len(group)]): venue_map[sid] = v

                df_target['場地'] = df_target['學號'].map(venue_map).fillna("")
                
                df_teacher['監考教師'] = get_str_col(df_teacher, ['監考教師', '監考老師', '教師姓名', '老師'])
                df_teacher['場地'] = get_str_col(df_teacher, ['場地', '地點', '考場', '場點'])
                df_teacher['比對年級'] = get_str_col(df_teacher, ['監考年級', '年級']).apply(grade_to_chinese)
                t_map = df_teacher.drop_duplicates(subset=['比對年級']).set_index('比對年級')[get_str_col(df_teacher, ['時間']).name].to_dict()
                df_target['時間2'] = df_target['年級'].map(t_map).fillna("")

                # --- 階段二：報表二處理 (標籤) ---
                df_label = df_target.drop_duplicates(subset=['學號', '試卷編號'], keep='first').copy()
                df_label['單科標籤'] = df_label.apply(lambda r: f"{r['試卷編號']}{r['科目簡稱']}" if str(r['試卷編號']).strip() != "" else "", axis=1)
                
                group_cols = ['年級', '班級', '座號', '姓名', '科目數目', '場地'] 
                df_grouped = df_label.groupby(group_cols, dropna=False, as_index=False).agg({
                    '單科標籤': lambda x: '、'.join(sorted(dict.fromkeys([str(i) for i in x if str(i).strip() != ""]), key=natural_sort_key))
                })
                
                df_grouped = df_grouped.rename(columns={'科目數目': '個人考科', '場地': '地點', '單科標籤': '所有考科'})
                
                df_vld = df_grouped[df_grouped['地點'] != ""].copy()
                df_vld['G_W'] = df_vld['年級'].map(grade_weight).fillna(99)
                df_vld['L_W'] = df_vld['地點'].map(loc_weight).fillna(99)
                df_vld['NumSeat'] = pd.to_numeric(df_vld['座號'], errors='coerce').fillna(999)
                df_vld = df_vld.sort_values(by=['G_W', 'L_W', '個人考科', 'NumSeat', '班級'], ascending=[True, True, False, True, True])
                
                f_dfs = []
                for gr in ['一', '二', '三']:
                    for loc, cap in current_targets.items():
                        sub = df_vld[(df_vld['地點'] == loc) & (df_vld['年級'] == gr)].copy()
                        if sub.empty: continue 
                        if len(sub) < cap: 
                            pad = pd.DataFrame([{'地點': loc, '年級': gr}] * (cap - len(sub)))
                            sub = pd.concat([sub, pad], ignore_index=True)
                        sub['序號'] = [f"{i+1:03d}" for i in range(len(sub))]
                        f_dfs.append(sub)
                
                label_cols = ['所有考科', '班級', '年級', '座號', '姓名', '個人考科', '地點', '序號']
                df_rep2 = pd.concat(f_dfs, ignore_index=True) if f_dfs else pd.DataFrame(columns=label_cols)
                for col in label_cols:
                    if col not in df_rep2.columns: df_rep2[col] = ""
                df_rep2 = df_rep2[label_cols]

                # --- 階段三：報表三處理 (考程) ---
                df_exam = df_target.drop_duplicates(subset=['學號', '試卷編號'], keep='first').copy()

                if '試卷編號' in df_exam.columns:
                    clean_papers = df_exam['試卷編號'].fillna('').astype(str).str.replace(r'\.0$', '', regex=True).str.replace(r'\s+', '', regex=True).str.upper()
                    df_exam = df_exam[~clean_papers.isin(['', 'NAN', 'NONE', '<NA>', 'NULL'])].copy()
                    df_exam['試卷編號'] = clean_papers[~clean_papers.isin(['', 'NAN', 'NONE', '<NA>', 'NULL'])]

                def extract_loc_short(loc):
                    l = str(loc); return '致用' if '致用' in l else '圖書' if '圖書' in l else '電腦' if '電腦' in l else l

                df_exam['比對場地'] = df_exam['場地'].apply(extract_loc_short) 
                df_teacher['比對場地'] = df_teacher['場地'].apply(extract_loc_short)
                df_exam['應到人數'] = df_exam.groupby(['場地', '班級', '科目簡稱'])['學號'].transform('count')
                
                valid_locs_exam = df_exam[df_exam['比對場地'].isin(['致用', '圖書', '電腦'])]
                loc_counts = valid_locs_exam.groupby(['班級', '試卷編號'])['比對場地'].nunique().to_dict()
                
                def determine_group_id(row):
                    c, p, l = str(row['班級']), str(row['試卷編號']), str(row['比對場地'])
                    if loc_counts.get((c, p), 0) <= 1: return p 
                    suffix = ".1" if l == '致用' else ".2" if l == '圖書' else ".3" if l == '電腦' else ""
                    return f"{p}{suffix}"
                    
                df_exam['分組編號'] = df_exam.apply(determine_group_id, axis=1)
                teacher_map = df_teacher.drop_duplicates(subset=['比對年級', '比對場地']).set_index(['比對年級', '比對場地'])['監考教師'].to_dict()
                df_exam['監考教師'] = df_exam.apply(lambda r: teacher_map.get((r['年級'], r['比對場地']), ""), axis=1)

                final_cols = ['班級', '座號', '學號', '姓名', '科目簡稱', '試卷編號', '分組編號', '場地', '授課教師', '監考教師', '時間2', '應到人數']
                for col in final_cols:
                    if col not in df_exam.columns: df_exam[col] = ""

                df_final_exam = df_exam[final_cols].copy()
                
                df_final_exam['G_W'] = df_final_exam['班級'].apply(grade_to_chinese).map(grade_weight).fillna(99)
                df_final_exam['L_W'] = df_final_exam['場地'].map(loc_weight).fillna(99)
                
                df_final_exam = df_final_exam.sort_values(
                    by=['G_W', 'L_W', '場地', '應到人數', '班級', '科目簡稱', '座號'], 
                    ascending=[True, True, True, True, True, True, True] 
                )

                df_final_exam['GroupKey'] = df_final_exam['場地'] + "_" + df_final_exam['班級'] + "_" + df_final_exam['科目簡稱']
                
                grouped = [g for _, g in df_final_exam.groupby('GroupKey', sort=False)]
                final_rows = []
                empty = pd.DataFrame([[np.nan] * len(final_cols)], columns=final_cols)
                
                for i, grp in enumerate(grouped):
                    final_rows.append(grp.drop(columns=['GroupKey', 'G_W', 'L_W']))
                    if i < len(grouped) - 1: final_rows.append(empty)
                
                if final_rows:
                    df_rep3_final = pd.concat(final_rows, ignore_index=True).fillna("")
                else:
                    df_rep3_final = pd.DataFrame(columns=final_cols)

                # --- 階段四：報表四處理 (印卷) ---
                df_rep4 = df_target[df_target['試卷編號'] != ""].drop_duplicates(subset=['學號', '試卷編號']).groupby('試卷編號').size().reset_index(name='試卷數量')
                df_rep4['SortKey'] = df_rep4['試卷編號'].apply(natural_sort_key)
                df_rep4 = df_rep4.sort_values(by='SortKey').drop(columns=['SortKey'])

                # --- 階段五：報表五處理 (全校補考範圍表) ---
                df_rep5 = df_target[df_target['科目'] != ""].drop_duplicates(subset=['班級', '科目']).copy()
                df_rep5['G_W'] = df_rep5['班級'].apply(grade_to_chinese).map(grade_weight).fillna(99)
                df_rep5 = df_rep5.sort_values(by=['G_W', '班級', '科目'])
                df_rep5 = df_rep5[['班級', '科目', '補考範圍']].rename(columns={'科目': '所有補考的科目'})

                df_rep5['補考範圍'] = df_rep5['補考範圍'].apply(
                    lambda x: "請詢問該科任課教師" if str(x).strip() in ["", "nan", "None", "NaN", "<NA>"] else str(x).strip()
                )

                if HAS_DOCX:
                    scope_bytes = to_word_scope_bytes(df_rep5)
                    scope_ext = "docx"
                    scope_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    scope_bytes = to_excel_scope_bytes(df_rep5)
                    scope_ext = "xlsx"
                    scope_mime = "application/vnd.ms-excel"

                # ==========================================
                # ⭐ 階段六：報表六處理 (補考學生名冊 - 排版版)
                # ==========================================
                df_rep6 = df_target[df_target['科目'] != ""].drop_duplicates(subset=['學號', '科目']).copy()
                
                # 套用姓名遮蔽防護
                df_rep6['姓名'] = df_rep6['姓名'].apply(mask_name_func)
                df_rep6 = df_rep6.rename(columns={'時間2': '補考時間', '場地': '補考地點'})
                
                # 針對沒有「試卷編號」的科目，修改時間與清空地點
                no_paper_mask = df_rep6['試卷編號'].astype(str).str.strip() == ""
                formatted_date = f"{self_exam_deadline.month}月{self_exam_deadline.day}日"
                custom_time_msg = f"請於 {formatted_date} 前自行找老師補考"
                
                df_rep6.loc[no_paper_mask, '補考時間'] = custom_time_msg
                df_rep6.loc[no_paper_mask, '補考地點'] = ""
                
                # 挑選指定欄位與排序
                cols_to_keep = ['班級', '座號', '學號', '姓名', '科目', '補考時間', '補考地點']
                for c in cols_to_keep:
                    if c not in df_rep6.columns: df_rep6[c] = ""
                df_rep6 = df_rep6[cols_to_keep]
                
                df_rep6['G_W'] = df_rep6['班級'].apply(grade_to_chinese).map(grade_weight).fillna(99)
                df_rep6['NumSeat'] = pd.to_numeric(df_rep6['座號'], errors='coerce').fillna(999)
                df_rep6 = df_rep6.sort_values(by=['G_W', '班級', 'NumSeat', '科目']).drop(columns=['G_W', 'NumSeat'])

                # ⭐ 套用自訂的表頭排版匯出 (帶入學年度與日期，單日雙日皆支援)
                student_list_bytes = to_excel_student_list_bytes(df_rep6, school_year_input, exam_date_1, exam_date_2)

                # --- 欄位重新命名以完美銜接合併列印 ---
                rename_mapping = {'姓名': '學生姓名', '場地': '地點'}
                df_target_out = df_target.rename(columns=rename_mapping)
                df_rep2_out = df_rep2.rename(columns=rename_mapping)
                df_rep3_out = df_rep3_final.rename(columns=rename_mapping)
                df_rep4_out = df_rep4.rename(columns=rename_mapping)

                # 【鎖定記憶體】
                st.session_state['results'] = {
                    'venue': to_excel_bytes(df_target_out),
                    'label': to_excel_bytes(df_rep2_out),
                    'schedule': to_excel_bytes(df_rep3_out),
                    'print': to_excel_bytes(df_rep4_out),
                    'scope': scope_bytes,
                    'scope_ext': scope_ext,
                    'scope_mime': scope_mime,
                    'student_list': student_list_bytes
                }
                st.balloons()

            except Exception as e:
                st.error("🚨 發生未預期錯誤，請檢查檔案格式是否正確。")
                with st.expander("點此查看詳細工程錯誤碼"):
                    st.code(traceback.format_exc())

# ==========================================
# 5. 下載區
# ==========================================
if st.session_state['results'] is not None:
    st.divider()
    st.success("🎊 運算結果已鎖定，您可以逐一下載所有 6 份檔案：")
    
    res = st.session_state['results']
    d_col1, d_col2 = st.columns(2)
    
    with d_col1:
        st.download_button("📄 下載：1.場地分配版", res['venue'], "1_場地分配版.xlsx", "application/vnd.ms-excel", use_container_width=True)
        st.download_button("📋 下載：3.考程匯整表", res['schedule'], "3_全校補考考程匯整表.xlsx", "application/vnd.ms-excel", use_container_width=True)
        scope_filename = f"5_全校補考範圍表_獨立公告版.{res['scope_ext']}"
        st.download_button(f"📖 下載：5.補考範圍表 ({res['scope_ext'].upper()}分頁)", res['scope'], scope_filename, res['scope_mime'], use_container_width=True)
        
    with d_col2:
        st.download_button("🖨️ 下載：2.排座標籤", res['label'], "2_報表二_排座標籤.xlsx", "application/vnd.ms-excel", use_container_width=True)
        st.download_button("📝 下載：4.試卷印製表", res['print'], "4_試卷印製數量表.xlsx", "application/vnd.ms-excel", use_container_width=True)
        st.download_button("🧑‍🎓 下載：6.補考學生名冊", res['student_list'], "6_補考學生名冊(含注意事項版).xlsx", "application/vnd.ms-excel", use_container_width=True)
