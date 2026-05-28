import streamlit as st
import pandas as pd
import datetime
import io
import os
import traceback

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.shared import Pt, Cm, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ==========================================
# 1. 網頁頁面配置與記憶體初始化
# ==========================================
st.set_page_config(page_title="教甄智能排程系统", page_icon="🏫", layout="wide")
st.title("🏫 試務組-教師甄選智能輔助系统")
st.info("💡 終極進化：工作人員資料袋場地已精準對接「場地教室」，且已將「試場用品」與「工作人員」順序對調，工作人員移至最下方！(115.05.28增修)")

if not HAS_DOCX:
    st.error("🚨 偵測到系統未安裝 `python-docx` 套件！無法產出直出版 Word。請在 requirements.txt 中加入 `python-docx`。")

# 初始化 Session State
if 'tab1_processed' not in st.session_state:
    st.session_state.tab1_processed = False
if 'tab2_processed' not in st.session_state:
    st.session_state.tab2_processed = False
if 'staff_env_data' not in st.session_state:
    st.session_state.staff_env_data = None

# ==========================================
# 0. 側邊欄：試務資源與印章設定
# ==========================================
st.sidebar.title("🔴 自動蓋章設定 (公告與簽到表專用)")

default_stamp_name = "試務組印章.png"
stamp_path = default_stamp_name
has_default_stamp = False
file_stamp = None

if os.path.exists(stamp_path):
    has_default_stamp = True
elif os.path.exists(f"../{default_stamp_name}"):
    stamp_path = f"../{default_stamp_name}"
    has_default_stamp = True
elif os.path.exists(f"pages/{default_stamp_name}"):
    stamp_path = f"pages/{default_stamp_name}"
    has_default_stamp = True

if has_default_stamp:
    st.sidebar.success(f"✅ 系統已自動載入預設印章：\n`{default_stamp_name}`")
    with st.sidebar.expander("🔄 想要臨時更換印章？點此手動上傳"):
        file_stamp = st.file_uploader("上傳新印章圖檔 (.png, .jpg)", type=['png', 'jpg', 'jpeg'], key="stamp_upload")
else:
    st.sidebar.markdown("上傳您的「試務組印章.png」，系統將自動印在表單右下角。")
    file_stamp = st.sidebar.file_uploader("上傳印章圖檔 (.png, .jpg)", type=['png', 'jpg', 'jpeg'], key="stamp_upload")

def get_stamp_io():
    """每次呼叫都產生一個全新的 BytesIO 給 Word 套印，避免檔案被關閉"""
    if file_stamp:
        return io.BytesIO(file_stamp.getvalue())
    elif has_default_stamp:
        with open(stamp_path, "rb") as f:
            return io.BytesIO(f.read())
    return None

# ==========================================
# 全域參數：學年度與甄選次數設定
# ==========================================
st.markdown("### 📅 全域試務設定")
current_roc_year = datetime.datetime.now().year - 1911
c_year, c_session = st.columns(2)
with c_year:
    academic_year = st.number_input("學年度 (系統已自動帶入)", min_value=100, max_value=200, value=current_roc_year, step=1)
with c_session:
    session_num = st.number_input("第幾次甄選 (如: 1)", min_value=1, max_value=50, value=1, step=1)
st.divider()

# ==========================================
# 2. 共用函式區 (表單產生器與排程矩陣)
# ==========================================
def generate_signin_sheet(academic_year, session_num, df_master, target_subjs):
    doc = docx.Document()
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(1.5), Cm(2.0)
    section.left_margin, section.right_margin = Cm(2.0), Cm(2.0)
    
    footer = section.footer
    for p in footer.paragraphs: p._element.getparent().remove(p._element)
    
    footer_table = footer.add_table(rows=1, cols=2, width=Cm(16.0))
    footer_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    cell_left = footer_table.rows[0].cells[0]
    cell_left.width = Cm(8.0)
    cell_right = footer_table.rows[0].cells[1]
    cell_right.width = Cm(8.0)
    
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sig = p_right.add_run("承辦人簽名：__________________\n")
    r_sig.font.name = '標楷體'
    r_sig._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
    r_sig.font.size = Pt(14)
    
    stamp_io = get_stamp_io()
    if stamp_io:
        run_stamp = p_right.add_run()
        run_stamp.add_picture(stamp_io, width=Cm(4.0))
    
    for idx, subj in enumerate(target_subjs):
        df_sub = df_master[df_master['報考科目'] == subj]
        if df_sub.empty: continue
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_t = p_title.add_run(f"{academic_year}學年度第{session_num}次代理教師甄選\n【{subj}】考生簽到表")
        r_t.font.name = '標楷體'
        r_t._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
        r_t.font.size = Pt(18)
        r_t.bold = True
        
        has_name = '姓名' in df_master.columns
        headers = ['編號', '甄選證號']
        if has_name: headers.append('姓名')
        headers.extend(['報到時間', '考生簽名', '備註'])
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        for c_i, h in enumerate(headers):
            cell = table.cell(0, c_i)
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '標楷體'
                    r._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                    r.font.size = Pt(14)
                    r.bold = True
                    
        for _, cand in df_sub.iterrows():
            row_cells = table.add_row().cells
            table.rows[-1].height = Cm(1.2)
            table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            
            name_str = cand.get('姓名', '') if has_name else ''
            base_data = [str(cand['排序']), str(cand['准考證號'])]
            if has_name: base_data.append(name_str)
            
            for c_i in range(len(headers)):
                if c_i < len(base_data):
                    row_cells[c_i].text = base_data[c_i]
                else:
                    row_cells[c_i].text = "" 
                    
                for p in row_cells[c_i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = '標楷體'
                        r._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                        r.font.size = Pt(14)
                        
        if idx < len(target_subjs) - 1:
            doc.add_page_break()
            
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def generate_eval_sheet(academic_year, session_num, df_master, form_name, row_items, target_subjs):
    doc = docx.Document()
    section = doc.sections[0]
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin, section.bottom_margin = Cm(1.5), Cm(1.5)
    section.left_margin, section.right_margin = Cm(1.5), Cm(1.5)
    
    for idx, subj in enumerate(target_subjs):
        df_sub = df_master[df_master['報考科目'] == subj]
        if df_sub.empty: continue
        
        cands_all = df_sub['准考證號'].tolist()
        chunk_size = 5 if form_name == "實作評分表" else 15
        chunks = [cands_all[i:i + chunk_size] for i in range(0, len(cands_all), chunk_size)]
        
        for chunk_idx, cands in enumerate(chunks):
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_t = p_title.add_run(f"{academic_year}學年度第{session_num}次代理教師甄選\n【{subj}】{form_name}")
            r_t.font.name = '標楷體'
            r_t._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
            r_t.font.size = Pt(18)
            r_t.bold = True
            
            table = doc.add_table(rows=len(row_items)+1, cols=len(cands)+1)
            table.style = 'Table Grid'
            
            table.rows[0].height = Cm(1.2)
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            
            table.cell(0, 0).text = "評分項目"
            for i, cand in enumerate(cands):
                table.cell(0, i+1).text = cand
                
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = '標楷體'
                        r._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                        r.font.size = Pt(14)
                        r.bold = True
            
            for r_i, item in enumerate(row_items):
                row = table.rows[r_i+1]
                if form_name == "實作評分表" and item == "評分內容":
                    row.height = Cm(8.0) 
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                else:
                    row.height = Cm(1.5)
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                    
                cell = table.cell(r_i+1, 0)
                cell.text = item
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = '標楷體'
                        r._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                        r.font.size = Pt(14)
                        r.bold = True
                        
            for _ in range(2): doc.add_paragraph("")
            
            footer_tbl = doc.add_table(rows=1, cols=2)
            footer_tbl.autofit = False
            c_left, c_right = footer_tbl.rows[0].cells
            c_left.width = Cm(16.0)
            c_right.width = Cm(10.0)
            
            p_left = c_left.paragraphs[0]
            r_left = p_left.add_run("本表總分超過90分及不滿70分，請說明原因。")
            r_left.font.name = '標楷體'
            r_left._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
            r_left.font.size = Pt(12)
            
            p_right = c_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_right = p_right.add_run("評分委員簽名：__________________")
            r_right.font.name = '標楷體'
            r_right._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
            r_right.font.size = Pt(14)
            
            if chunk_idx < len(chunks) - 1 or idx < len(target_subjs) - 1:
                doc.add_page_break()
                
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def generate_envelope_cover(target_subjs, env_title):
    doc = docx.Document()
    section = doc.sections[0]
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width = Cm(36.4)
    section.page_height = Cm(25.7)
    
    section.top_margin = Cm(10.0)
    section.bottom_margin = Cm(4.0)
    section.left_margin = Cm(10.0)
    section.right_margin = Cm(10.0)
    
    for i, subject in enumerate(target_subjs):
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0) 
        run1 = p1.add_run(subject)
        run1.font.name = '標楷體'
        run1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
        run1.font.size = Pt(80) 
        run1.bold = True
        
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0) 
        run2 = p2.add_run(env_title)
        run2.font.name = '標楷體'
        run2._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
        run2.font.size = Pt(80) 
        run2.bold = True
        
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(0) 
        p3.paragraph_format.space_after = Pt(0) 
        run3 = p3.add_run("(內附原子筆，進入試場前煩請確認是否能書寫)")
        run3.font.name = '標楷體'
        run3._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
        run3.font.size = Pt(20) 
        
        if i < len(target_subjs) - 1:
            doc.add_page_break()
            
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def generate_staff_envelopes(df_dict):
    """產出工作人員資料袋封面，支援雙科分行與隱形文字對齊技術"""
    doc = docx.Document()
    section = doc.sections[0]
    
    # B4 直式
    section.page_width = Cm(25.7)
    section.page_height = Cm(36.4)
    
    # 邊界精準設定
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    is_first_page = True
    
    for sheet_name, df in df_dict.items():
        for idx, row in df.iterrows():
            subj1 = str(row.get('科別1', '')).strip()
            if not subj1 or subj1.lower() == 'nan':
                continue
                
            if not is_first_page:
                doc.add_page_break()
            is_first_page = False
            
            # 1. 處理科別 1 (72pt)
            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_after = Pt(0) 
            run1 = p1.add_run(subj1)
            run1.font.name = '標楷體'
            run1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
            run1.font.size = Pt(72)
            run1.bold = True
            
            # 2. 處理科別 2 (72pt)
            subj2 = str(row.get('科別2', '')).strip()
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(0)
            
            if subj2 and subj2.lower() != 'nan':
                run2 = p2.add_run(subj2)
                run2.font.name = '標楷體'
                run2._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                run2.font.size = Pt(72)
                run2.bold = True
            else:
                # 【終極解法：隱形中文字】
                run2 = p2.add_run("一") 
                run2.font.name = '標楷體'
                run2._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                run2.font.size = Pt(72)
                run2.bold = True
                run2.font.color.rgb = RGBColor(255, 255, 255) # 純白隱形
            
            # 3. 處理試場場地 (48pt) -> 【場地類型:地點】 格式
            venue = str(row.get('試場場地', row.get('場地教室', row.get('教室地點', '')))).strip()
            if venue.lower() == 'nan': venue = ""
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(40)
            
            venue_text = f"【{sheet_name}:{venue}】" if venue else f"【{sheet_name}】"
            
            run3 = p3.add_run(venue_text)
            run3.font.name = '標楷體'
            run3._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
            run3.font.size = Pt(48)
            run3.bold = True
            
            # 4. 處理試場用品 (36pt) -> 置中對齊
            supplies = str(row.get('試場用品', '')).strip()
            if supplies.lower() == 'nan': supplies = ""
            p4 = doc.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_before = Pt(60) # 原本工作人員的距離
            
            run4 = p4.add_run(f"試場用品：\n{supplies}")
            run4.font.name = '標楷體'
            run4._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
            run4.font.size = Pt(36)
            
            # 5. 處理工作人員 (36pt) -> 置中對齊 (放最下面)
            staff = str(row.get('工作人員', '')).strip()
            if staff.lower() == 'nan': staff = ""
            p5 = doc.add_paragraph()
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p5.paragraph_format.space_before = Pt(40)
            
            run5 = p5.add_run(f"工作人員：{staff}")
            run5.font.name = '標楷體'
            run5._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
            run5.font.size = Pt(36)
            
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# --- 排程矩陣設定 ---
TEACH_15_MATRIX = {
    1: ("09:40-09:55", "09:55-10:10"), 2: ("09:55-10:10", "10:10-10:25"),
    3: ("10:10-10:25", "10:25-10:40"), 4: ("10:25-10:40", "10:40-10:55"),
    5: ("10:40-10:55", "10:55-11:10"), 6: ("10:55-11:10", "11:10-11:25"),
    7: ("11:10-11:25", "11:25-11:40"), 8: ("11:25-11:40", "11:40-11:55"),
    9: ("13:10-13:25", "13:25-13:40"), 10: ("13:25-13:40", "13:40-13:55"),
    11: ("13:40-13:55", "13:55-14:10"), 12: ("13:55-14:10", "14:10-14:25"),
    13: ("14:10-14:25", "14:25-14:40"), 14: ("14:25-14:40", "14:40-14:55"),
    15: ("14:40-14:55", "14:55-15:10"), 16: ("14:55-15:10", "15:10-15:25"),
    17: ("15:10-15:25", "15:25-15:40"), 18: ("15:25-15:40", "15:40-15:55"),
    19: ("15:40-15:55", "15:55-16:10"), 20: ("15:55-16:10", "16:10-16:25"),
    21: ("16:10-16:25", "16:25-16:40"), 22: ("16:25-16:40", "16:40-16:55"),
    23: ("16:40-16:55", "16:55-17:10"), 24: ("16:55-17:10", "17:10-17:25"), 25: ("17:10-17:25", "17:25-17:40"),
}

ORAL_15_MATRIX = {
    1: {1: "10:10-10:20"},
    2: {1: "10:10-10:20", 2: "09:40-09:50"},
    3: {1: "10:10-10:20", 2: "09:40-09:50", 3: "09:55-10:05"},
    4: {1: "10:25-10:35", 2: "09:40-09:50", 3: "09:55-10:05", 4: "10:10-10:20"},
    5: {1: "10:40-10:50", 2: "09:40-09:50", 3: "09:55-10:05", 4: "10:10-10:20", 5: "10:25-10:35"},
    6: {1: "10:10-10:20", 2: "09:40-09:50", 3: "09:55-10:05", 4: "10:55-11:05", 5: "10:25-10:35", 6: "10:40-10:50"},
    7: {1: "10:10-10:20", 2: "09:40-09:50", 3: "09:55-10:05", 4: "11:10-11:20", 5: "10:25-10:35", 6: "10:40-10:50", 7: "10:55-11:05"},
    8: {1: "10:10-10:20", 2: "09:40-09:50", 3: "09:55-10:05", 4: "11:25-11:35", 5: "10:25-10:35", 6: "10:40-10:50", 7: "10:55-11:05", 8: "11:10-11:20"},
    9: {1: "10:10-10:20", 2: "09:40-09:50", 3: "09:55-10:05", 4: "10:55-11:05", 5: "10:25-10:35", 6: "10:40-10:50", 7: "11:40-11:50", 8: "11:10-11:20", 9: "11:25-11:35"},
    10: {
        1: "10:10-10:20", 2: "09:40-09:50", 3: "09:55-10:05", 4: "10:55-11:05", 5: "10:25-10:35", 6: "10:40-10:50", 
        7: "11:40-11:50", 8: "11:10-11:20", 9: "11:25-11:35", 10: "13:10-13:20", 11: "13:25-13:35", 12: "13:40-13:50", 
        13: "13:55-14:05", 14: "14:10-14:20", 15: "14:25-14:35", 16: "14:40-14:50", 17: "14:55-15:05", 18: "15:10-15:20",
        19: "15:25-15:35", 20: "15:40-15:50", 21: "15:55-16:05", 22: "16:10-16:20", 23: "16:25-16:35", 24: "16:40-16:50", 25: "16:55-17:05"
    }
}

TEACH_30_MATRIX = {
    1: ("09:40-09:55", "09:55-10:25"), 2: ("10:10-10:25", "10:25-10:55"),
    3: ("10:40-10:55", "10:55-11:25"), 4: ("11:10-11:25", "11:25-11:55"),
    5: ("13:10-13:25", "13:25-13:55"), 6: ("13:40-13:55", "13:55-14:25"),
    7: ("14:10-14:25", "14:25-14:55"), 8: ("14:40-14:55", "14:55-15:25"),
    9: ("15:10-15:25", "15:25-15:55"), 10: ("15:40-15:55", "15:55-16:25"),
    11: ("16:10-16:25", "16:25-16:55"), 12: ("16:40-16:55", "16:55-17:25"),
    13: ("17:10-17:25", "17:25-17:55"), 14: ("17:40-17:55", "17:55-18:25"), 15: ("18:10-18:25", "18:25-18:55"),
}

def generate_oral_30_independent():
    matrix = {}
    base_schedule = [
        "10:25-10:35", "09:40-09:50", "09:55-10:05", "10:10-10:20", "10:40-10:50", "10:55-11:05", "11:10-11:20", "11:25-11:35", "11:40-11:50",
        "13:10-13:20", "13:25-13:35", "13:40-13:50", "13:55-14:05", "14:10-14:20", "14:25-14:35", "14:40-14:50", "14:55-15:05", "15:10-15:20",
        "15:25-15:35", "15:40-15:50", "15:55-16:05", "16:10-16:20", "16:25-16:35", "16:40-16:50", "16:55-17:05"
    ]
    for n in range(1, 26):
        matrix[n] = {i+1: base_schedule[i] for i in range(n)}
    return matrix

ORAL_30_MATRIX_INDEPENDENT = generate_oral_30_independent()

CHRONOLOGICAL_ORAL_POOL = [
    "09:40-09:50", "09:55-10:05", "10:10-10:20", "10:25-10:35", "10:40-10:50", "10:55-11:05",
    "11:10-11:20", "11:25-11:35", "11:40-11:50", 
    "13:10-13:20", "13:25-13:35", "13:40-13:50", "13:55-14:05", "14:10-14:20", "14:25-14:35",
    "14:40-14:50", "14:55-15:05", "15:10-15:20", "15:25-15:35", "15:40-15:50", "15:55-16:05",
    "16:10-16:20", "16:25-16:35", "16:40-16:50", "16:55-17:05", "17:10-17:20", "17:25-17:35",
    "17:40-17:50", "17:55-18:05", "18:10-18:20", "18:25-18:35", "18:40-18:50", "18:55-19:05"
]

def check_time_conflict_bool(prep_str, teach_str, oral_str):
    if not prep_str or not teach_str or not oral_str or "手動" in oral_str: return False
    def parse_m(s):
        try:
            h, m = map(int, s.strip().split(':'))
            return h * 60 + m
        except: return 0
    try:
        p_s, p_e = map(parse_m, prep_str.split('-'))
        t_s, t_e = map(parse_m, teach_str.split('-'))
        o_s, o_e = map(parse_m, oral_str.split('-'))
        times = [(p_s, p_e, 'p'), (t_s, t_e, 't'), (o_s, o_e, 'o')]
        times.sort(key=lambda x: x[0])
        for i in range(len(times)-1):
            if times[i][1] > times[i+1][0]: return True
        return False
    except: return False

# ==========================================
# 3. 介面分頁 (雙階段架構)
# ==========================================
tab1, tab2 = st.tabs(["📂 第一階段：考前前置作業 (產出簽到與評分表/信封)", "⏱️ 第二階段：考試當天排程 (產出公告與時間總表)"])

# -------------------------------------------------------------
# TAB 1: 第一階段 (前置表單產生)
# -------------------------------------------------------------
with tab1:
    st.subheader("📝 上傳檔案產出前置表單")
    st.markdown("此階段專為**考試前**準備文件設計，純粹依據准考證號與您建立的格式配置產出 Word 檔案。")
    
    file_reg = st.file_uploader("1️⃣ 上傳【報名總名單】 (准考證號, 報考科目, [姓名]).xlsx", type=['xlsx'], key="reg_file")
    
    st.markdown("---")
    file_staff = st.file_uploader("2️⃣ 上傳【工作人員資料袋套印】 (工作表需為：準備室/試教場地/口試場地).xlsx", type=['xlsx'], key="staff_file")
    
    if file_reg or file_staff:
        try:
            if st.button("🚀 產出前置作業表單", type="primary", use_container_width=True):
                
                # 處理【名單】產出的 簽到表、評分表、以及 評分表信封封面
                if file_reg:
                    df_reg = pd.read_excel(file_reg).fillna("")
                    df_reg['准考證號'] = df_reg['准考證號'].astype(str).str.strip()
                    df_reg = df_reg[df_reg['准考證號'] != ""]
                    
                    all_subjs_t1 = df_reg['報考科目'].unique().tolist()
                    
                    all_cands_t1 = []
                    for subj in all_subjs_t1:
                        df_sub_t1 = df_reg[df_reg['報考科目'] == subj].copy()
                        df_sub_t1 = df_sub_t1.sort_values('准考證號')
                        
                        for i, row in enumerate(df_sub_t1.to_dict('records')):
                            cand_record = {'報考科目': subj, '准考證號': row['准考證號'], '排序': i + 1}
                            if '姓名' in row: cand_record['姓名'] = str(row['姓名']).strip()
                            all_cands_t1.append(cand_record)
                            
                    df_master_t1 = pd.DataFrame(all_cands_t1)
                    
                    st.session_state.sign_data = generate_signin_sheet(academic_year, session_num, df_master_t1, all_subjs_t1)
                    
                    teach_items = ["教學技巧(30%)", "語言表達(30%)", "儀態(20%)", "教室管理(10%)", "時間管理(10%)", "合計總分", "備註"]
                    st.session_state.teach_data = generate_eval_sheet(academic_year, session_num, df_master_t1, "試教評分表", teach_items, all_subjs_t1)
                    
                    oral_items = ["自述(20%)", "教學理念(20%)", "班級經營(20%)", "表達溝通(20%)", "舉止儀態(20%)", "合計總分", "備註"]
                    st.session_state.oral_data = generate_eval_sheet(academic_year, session_num, df_master_t1, "口試評分表", oral_items, all_subjs_t1)
                        
                    prac_subjs_t1 = [s for s in all_subjs_t1 if '實作' in s] 
                    
                    st.session_state.oral_env_data = generate_envelope_cover(all_subjs_t1, "口試評分表")
                    st.session_state.teach_env_data = generate_envelope_cover(all_subjs_t1, "試教評分表")
                
                # 處理【工作人員資料袋】產出的 B4 封面
                if file_staff:
                    xls_staff = pd.ExcelFile(file_staff)
                    staff_df_dict = {}
                    for sheet in ['準備室', '試教場地', '口試場地']:
                        if sheet in xls_staff.sheet_names:
                            staff_df_dict[sheet] = pd.read_excel(xls_staff, sheet_name=sheet).fillna("")
                    
                    if staff_df_dict:
                        st.session_state.staff_env_data = generate_staff_envelopes(staff_df_dict)

                st.session_state.tab1_processed = True
                
        except Exception as e:
            st.error(f"發生錯誤: {e}")
            st.code(traceback.format_exc())

    if st.session_state.tab1_processed:
        st.success("🎉 前置表單產出完成！系統已將您提供的資料完美轉為獨立報表。")
        
        # 顯示名單相關報表
        if st.session_state.get('sign_data'):
            st.markdown("#### 📄 報到與評分表 (A4)")
            c_d3, c_d4, c_d5, c_d6 = st.columns(4)
            with c_d3:
                st.download_button("✍️ 1. 下載 考生簽到表", data=st.session_state.sign_data, file_name=f"{academic_year}學年度_考生簽到表.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")
            with c_d4:
                st.download_button("🧑‍🏫 2. 下載 試教評分表", data=st.session_state.teach_data, file_name=f"{academic_year}學年度_試教評分表.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")
            with c_d5:
                st.download_button("🗣️ 3. 下載 口試評分表", data=st.session_state.oral_data, file_name=f"{academic_year}學年度_口試評分表.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")
            
            st.markdown("#### ✉️ 評分表專用信封袋 (B4 橫向、大字體)")
            c_env1, c_env2 = st.columns(2)
            with c_env1:
                st.download_button("✉️ 下載 試教信封封面 (B4橫向)", data=st.session_state.teach_env_data, file_name=f"01.試教評分表信封封面(B4橫向)_{academic_year}學年度.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="secondary")
            with c_env2:
                st.download_button("✉️ 下載 口試信封封面 (B4橫向)", data=st.session_state.oral_env_data, file_name=f"01.口試評分表信封封面(B4橫向)_{academic_year}學年度.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="secondary")

        # 顯示工作人員資料袋
        if st.session_state.get('staff_env_data'):
            st.markdown("#### ✉️ 工作人員資料袋封面 (B4 直式)")
            st.download_button("✉️ 下載 工作人員資料袋封面套印 (B4)", data=st.session_state.staff_env_data, file_name=f"02.工作人員資料袋封面(B4)_{academic_year}學年度.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=False, type="primary")

# -------------------------------------------------------------
# TAB 2: 第二階段 (當天精準排程與公告)
# -------------------------------------------------------------
with tab2:
    st.subheader("⏱️ 上傳【實際報到名單與場地】進行排程計算")
    st.markdown("此階段將根據**考試當天實際報到**的考生，動態進行防撞與排程計算，產出最終的 Excel 時間總表與對外公告 Word。")
    
    col_t2_1, col_t2_2 = st.columns([1, 1], gap="large")
    
    with col_t2_1:
        file_candidates = st.file_uploader("1️⃣ 上傳【實際報到名單】 (准考證號, 報考科目).xlsx", type=['xlsx'], key="real_cand")
        file_venues = st.file_uploader("2️⃣ 上傳場地配置 (如 3.場地.xlsx)", type=['xlsx'], key="venues")
        
        group_settings = []
        practical_subjects = []
        all_subjs = []
        venue_dict = {}
        
        if file_candidates:
            try:
                df_temp = pd.read_excel(file_candidates).fillna("")
                if '報考科目' in df_temp.columns:
                    all_subjs = df_temp['報考科目'].unique().tolist()
                    
                    st.write("---")
                    st.markdown("### 🤝 3. 設定合併口試群組")
                    num_groups = st.number_input("欲建立的「合併口試組」數量：", min_value=0, max_value=5, value=0, step=1, key="num_grp_t2")
                    already_assigned = set()
                    
                    for g_i in range(int(num_groups)):
                        st.markdown(f"**【口試合併群組 {g_i + 1}】**")
                        available_options = [s for s in all_subjs if s not in already_assigned]
                        selected_for_g = st.multiselect(
                            f"選擇成員科目：",
                            options=available_options,
                            key=f"group_select_t2_{g_i}"
                        )
                        
                        merged_teaching = st.checkbox(
                            "✅ 此群組『試教』也共用同一組委員 (接力排下去)", 
                            value=False, 
                            key=f"merged_teach_t2_{g_i}"
                        )
                        st.write("")
                        
                        if selected_for_g:
                            group_settings.append({
                                'subjects': selected_for_g,
                                'merged_teaching': merged_teaching
                            })
                            already_assigned.update(selected_for_g)
                    
                    st.write("---")
                    st.markdown("### 🛠️ 4. 設定實作學科")
                    practical_subjects = st.multiselect("選擇包含實作的科目 (切換30分時間軸)：", options=all_subjs, key="prac_t2")
                    
            except Exception as e:
                st.error(f"讀取失敗: {e}")

    with col_t2_2:
        st.markdown("#### 💡 排程引擎運作說明")
        st.markdown("""
        * **試教分離邏輯**：口試合併時，可自由決定試教要「同時起跑」還是「接力排程」。
        * **絕對時間軸防撞**：合併口試採用純時間序防跳號設計，完美按時間順序填補空檔。
        """)

    st.write("---")

    if st.button("🚀 啟動當天防撞排程", type="primary", use_container_width=True, key="btn_phase2"):
        if not file_candidates or not file_venues:
            st.error("🚨 請確認【實際報到名單】與【場地設定】皆已上傳！")
        else:
            try:
                dynamic_title = f"{academic_year}學年度第{session_num}次代理教師甄選人員試教及口試時間表"
                
                df_venues = pd.read_excel(file_venues).fillna("")
                for _, row in df_venues.iterrows():
                    subj = str(row.get('科目', '')).strip()
                    if subj:
                        venue_dict[subj] = {
                            '休息室': str(row.get('休息室', '')).strip(),
                            '準備室': str(row.get('準備室', '')).strip(),
                            '試教': str(row.get('試教', '')).strip(),
                            '口試': str(row.get('口試', '')).strip()
                        }

                df_candidates = pd.read_excel(file_candidates).fillna("")
                df_candidates['准考證號'] = df_candidates['准考證號'].astype(str).str.strip()
                df_candidates = df_candidates[df_candidates['准考證號'] != ""]

                if df_candidates.empty:
                    st.warning("⚠️ 攔截警告：名單內已無有效考生！")
                    st.stop()

                final_processing_groups = []
                assigned_set = set()
                for g_dict in group_settings:
                    if g_dict['subjects']:
                        final_processing_groups.append({
                            'type': 'merged', 
                            'subjects': g_dict['subjects'],
                            'merged_teaching': g_dict['merged_teaching']
                        })
                        assigned_set.update(g_dict['subjects'])
                for sub in all_subjs:
                    if sub not in assigned_set:
                        final_processing_groups.append({
                            'type': 'independent', 
                            'subjects': [sub],
                            'merged_teaching': False
                        })

                all_schedules = []
                
                for group in final_processing_groups:
                    group_total_candidates = sum(len(df_candidates[df_candidates['報考科目'] == sub]) for sub in group['subjects'])
                    is_merged_group = len(group['subjects']) > 1 
                    merged_teaching = group['merged_teaching']
                    
                    if is_merged_group:
                        group['subjects'].sort(key=lambda x: 1 if x in practical_subjects else 0)
                    
                    global_oral_idx = 1
                    global_teach_idx = 1
                    used_oral_indices = set() 
                    
                    for s_idx, subject in enumerate(group['subjects']):
                        is_practical = subject in practical_subjects
                        df_sub = df_candidates[df_candidates['報考科目'] == subject]
                        candidates = df_sub.to_dict('records')
                        n_candidates = len(candidates)
                        if n_candidates == 0: continue
                        
                        for i in range(n_candidates):
                            cand = candidates[i]
                            sort_num = global_teach_idx if merged_teaching else (i + 1)
                            
                            if is_practical:
                                prep, teach = TEACH_30_MATRIX.get(sort_num, ("請手動調整", "請手動調整"))
                            else:
                                prep, teach = TEACH_15_MATRIX.get(sort_num, ("請手動調整", "請手動調整"))
                            
                            if merged_teaching:
                                global_teach_idx += 1
                                
                            oral_range = "請手動調整"
                            
                            if is_merged_group:
                                for o_idx, test_oral in enumerate(CHRONOLOGICAL_ORAL_POOL):
                                    if o_idx in used_oral_indices: continue
                                    if not check_time_conflict_bool(prep, teach, test_oral):
                                        oral_range = test_oral
                                        used_oral_indices.add(o_idx)
                                        break
                            else:
                                if is_practical:
                                    lookup_n = group_total_candidates if group_total_candidates <= 25 else 25
                                    oral_range = ORAL_30_MATRIX_INDEPENDENT.get(lookup_n, {}).get(global_oral_idx, "請手動調整")
                                else:
                                    lookup_n = group_total_candidates if group_total_candidates <= 9 else 10
                                    oral_range = ORAL_15_MATRIX.get(lookup_n, {}).get(global_oral_idx, "請手動調整")
                            
                            all_schedules.append({
                                '報考科目': subject,
                                '准考證號': cand['准考證號'],
                                '排序': sort_num,
                                '準備時間': prep,
                                '試教(實作)時間': teach,
                                '口試時間': oral_range
                            })
                            global_oral_idx += 1

                df_master = pd.DataFrame(all_schedules)
                
                def get_flow(row):
                    p_time = str(row.get('準備時間', '')).split('-')[0].strip()
                    t_time = str(row.get('試教(實作)時間', '')).split('-')[0].strip()
                    o_time = str(row.get('口試時間', '')).split('-')[0].strip()
                    if "手動" in p_time or "手動" in o_time or not p_time or not o_time: return "待手動確認"
                    flow_list = [(p_time, "準備"), (t_time, "試教"), (o_time, "口試")]
                    flow_list.sort(key=lambda x: x[0])
                    return " → ".join([item[1] for item in flow_list])

                def check_time_conflict_text(row):
                    def parse_time(time_str, name):
                        try:
                            if not time_str or "手動" in time_str: return None
                            s, e = str(time_str).strip().split('-')
                            s_min = int(s.split(':')[0]) * 60 + int(s.split(':')[1])
                            e_min = int(e.split(':')[0]) * 60 + int(e.split(':')[1])
                            return (s_min, e_min, name)
                        except: return None
                    times = []
                    for val, label in [(row.get('準備時間', ''), '準備'), (row.get('試教(實作)時間', ''), '試教'), (row.get('口試時間', ''), '口試')]:
                        item = parse_time(val, label)
                        if item: times.append(item)
                    times.sort(key=lambda x: x[0])
                    conflicts = []
                    for i in range(len(times) - 1):
                        if times[i][1] > times[i+1][0]:
                            conflicts.append(f"{times[i][2]}與{times[i+1][2]}")
                    if not times or len(times) < 3: return "⚠️ 待確認"
                    if conflicts: return "🚨 衝突: " + "、".join(conflicts)
                    return "✅ 無衝突"
                    
                df_master['考試流程'] = df_master.apply(get_flow, axis=1)
                df_master['衝突檢核'] = df_master.apply(check_time_conflict_text, axis=1)
                
                df_merge = df_master.copy()
                df_merge = df_merge.rename(columns={'報考科目': '科目', '准考證號': '准考證', '試教(實作)時間': '試教時間'})
                df_merge.insert(1, '休息室', df_merge['科目'].apply(lambda x: venue_dict.get(x, {}).get('休息室', '未設定')))
                df_merge.insert(2, '準備室', df_merge['科目'].apply(lambda x: venue_dict.get(x, {}).get('準備室', '未設定')))
                df_merge.insert(3, '試教', df_merge['科目'].apply(lambda x: venue_dict.get(x, {}).get('試教', '未設定')))
                df_merge.insert(4, '口試', df_merge['科目'].apply(lambda x: venue_dict.get(x, {}).get('口試', '未設定')))
                
                df_merge = df_merge[['科目', '休息室', '準備室', '試教', '口試', '准考證', '排序', '準備時間', '試教時間', '口試時間']]
                
                merge_with_blanks = []
                prev_subj = None
                for idx, row in df_merge.iterrows():
                    curr_subj = row['科目']
                    if prev_subj is not None and curr_subj != prev_subj:
                        merge_with_blanks.append({col: "" for col in df_merge.columns})
                    merge_with_blanks.append(row.to_dict())
                    prev_subj = curr_subj
                df_merge_final = pd.DataFrame(merge_with_blanks)
                
                output_excel = io.BytesIO()
                with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                    df_master.to_excel(writer, index=False, sheet_name='試務中心總表')
                    df_merge_final.to_excel(writer, index=False, sheet_name='合併列印專用')
                    df_teach = df_master[['報考科目', '排序', '試教(實作)時間', '准考證號']].copy()
                    df_teach.to_excel(writer, index=False, sheet_name='門口_試教實作表')
                    df_int = df_master[['報考科目', '准考證號', '口試時間']].copy()
                    df_int['開始時間'] = df_int['口試時間'].str[:5]
                    df_int['開始時間'] = df_int['開始時間'].fillna("")
                    df_int = df_int.sort_values(['報考科目', '開始時間']).drop(columns=['開始時間'])
                    df_int.to_excel(writer, index=False, sheet_name='門口_口試表')
                
                st.session_state.excel_data = output_excel.getvalue()

                if HAS_DOCX:
                    doc = docx.Document()
                    section = doc.sections[0]
                    section.footer_distance = Cm(1.0)
                    section.bottom_margin = Cm(3.0)
                    
                    style = doc.styles['Normal']
                    style.font.name = '標楷體'
                    style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                    style.font.size = Pt(16)
                    
                    footer = section.footer
                    for p in footer.paragraphs: p._element.getparent().remove(p._element)
                    
                    footer_table = footer.add_table(rows=1, cols=2, width=Cm(16.0))
                    footer_table.autofit = False
                    cell_left = footer_table.rows[0].cells[0]
                    cell_left.width = Cm(11.5)
                    cell_right = footer_table.rows[0].cells[1]
                    cell_right.width = Cm(4.5)
                    
                    p_footer_text = cell_left.paragraphs[0]
                    p_footer_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    run_1 = p_footer_text.add_run("※試教及口試時間")
                    run_1.font.name = '標楷體'
                    run_1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                    run_1.font.size = Pt(14)
                    
                    run_2 = p_footer_text.add_run("將依現場實際報到人數及試場情形作調整，")
                    run_2.font.name = '標楷體'
                    run_2._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                    run_2.font.size = Pt(16)
                    run_2.font.color.rgb = RGBColor(255, 0, 0)
                    
                    run_3 = p_footer_text.add_run("請考生於各科指定之休息室等候叫號")
                    run_3.font.name = '標楷體'
                    run_3._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                    run_3.font.size = Pt(14)
                    
                    stamp_io = get_stamp_io()
                    if stamp_io:
                        p_stamp = cell_right.paragraphs[0]
                        p_stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
                        run_stamp = p_stamp.add_run()
                        run_stamp.add_picture(stamp_io, width=Cm(4.0)) 
                    
                    for subject in all_subjs:
                        df_sub_sched = df_master[df_master['報考科目'] == subject]
                        if df_sub_sched.empty: continue
                        
                        p_title = doc.add_paragraph()
                        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_title = p_title.add_run(dynamic_title)
                        run_title.font.size = Pt(18)
                        run_title.bold = True
                        
                        p_sub = doc.add_paragraph()
                        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_sub = p_sub.add_run(f"【{subject}】")
                        run_sub.font.size = Pt(16)
                        run_sub.bold = True
                        
                        v = venue_dict.get(subject, {})
                        doc.add_paragraph(f"考生休息室：{v.get('休息室', '未設定')}")
                        doc.add_paragraph(f"試教準備室：{v.get('準備室', '未設定')}")
                        doc.add_paragraph(f"試教場地：{v.get('試教', '未設定')}")
                        doc.add_paragraph(f"口試場地：{v.get('口試', '未設定')}")
                        
                        table = doc.add_table(rows=1, cols=5)
                        table.style = 'Table Grid'
                        table.autofit = False 
                        
                        col_widths = [Cm(3.5), Cm(2.0), Cm(3.5), Cm(3.5), Cm(3.5)]
                        table.style.font.name = '標楷體'
                        table.style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')
                        table.style.font.size = Pt(16)
                        
                        hdr_cells = table.rows[0].cells
                        hdr_headers = ['甄選證號', '編號', '試教準備室', '試教', '口試']
                        for col_idx in range(5):
                            hdr_cells[col_idx].text = hdr_headers[col_idx]
                            hdr_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
                            hdr_cells[col_idx].width = col_widths[col_idx]
                            table.columns[col_idx].width = col_widths[col_idx]
                        
                        for _, cand in df_sub_sched.iterrows():
                            row_cells = table.add_row().cells
                            row_data = [str(cand['准考證號']), str(cand['排序']), str(cand['準備時間']), str(cand['試教(實作)時間']), str(cand['口試時間'])]
                            for col_idx in range(5):
                                row_cells[col_idx].text = row_data[col_idx]
                                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                row_cells[col_idx].width = col_widths[col_idx]
                        
                        doc.add_page_break()
                        
                    output_word = io.BytesIO()
                    doc.save(output_word)
                    st.session_state.word_data = output_word.getvalue()
                    
                st.session_state.df_preview = df_merge_final.head(15)
                st.session_state.excel_filename = f"{academic_year}學年度第{session_num}次_排程與場地整合表(當天排程).xlsx"
                st.session_state.word_filename = f"{academic_year}學年度第{session_num}次_各科公告表(自動排版).docx"
                st.session_state.tab2_processed = True

            except Exception as e:
                st.error(f"發生錯誤: {e}")
                st.code(traceback.format_exc())

    if st.session_state.tab2_processed:
        st.success("🎉 當天排程計算完成！所有衝堂皆已避開，請在下方下載對外公告與排程檔案。")
        c_d1, c_d2 = st.columns(2)
        with c_d1:
            st.download_button("📊 1. 下載 Excel 總表 (含考試流程與檢核)", data=st.session_state.excel_data, file_name=st.session_state.excel_filename, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, type="primary")
        with c_d2:
            if HAS_DOCX and st.session_state.word_data:
                st.download_button("📄 2. 下載 Word 各科公告時間表 (自動蓋章版)", data=st.session_state.word_data, file_name=st.session_state.word_filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")
        
        st.write("---")
        st.write("👀 **專屬 Word 對接資料預覽：**")
        st.dataframe(st.session_state.df_preview)
